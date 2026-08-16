#!/usr/bin/env python3
"""Create the submission-audited Frontiers Original Research manuscript.

The original DOCX is never modified. Quantitative GNN values are read from the
fresh frozen result artifacts so the text and Table 1 cannot silently diverge.
"""
from __future__ import annotations

import json
import re
import shutil
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "data/results_v2/PrecisionPhage_Frontiers_manuscript.docx"
OUTDIR = ROOT / "submission"
OUTPUT = OUTDIR / "PrecisionPhage_Frontiers_Original_Research_revised.docx"
RD = ROOT / "data/results_v2"


def set_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)


def insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    new = Paragraph(node, paragraph._parent)
    if style:
        new.style = style
    new.add_run(text)
    return new


def mark_header_row(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def fmt_ci(row) -> str:
    return f"{row.auc:.3f} ({row.auc_lo:.3f}–{row.auc_hi:.3f})"


def replace_embedded_figures(docx_path: Path) -> None:
    replacements = {
        "word/media/image2.png": RD / "figure_main.png",
        "word/media/image3.png": RD / "fig_feature_importance.png",
        "word/media/image4.png": RD / "fig_phist_compare.png",
        "word/media/image5.png": RD / "cocktail_coverage.png",
        "word/media/image6.png": RD / "temporal_dynamics.png",
    }
    alt_texts = [
        ("Figure 1", "Three-panel leakage-controlled phage–host prediction summary. "
         "Gradient-boosted tree AUROC exceeds GraphSAGE across four holdout regimes; "
         "the middle panel reports the taxonomic and sequence-cluster leakage hierarchy; "
         "the ablation panel shows positive message-passing gains, largest in dual cold start."),
        ("Figure 2", "Horizontal bar charts of gradient-boosted tree gain importance. "
         "Phage and host genome composition dominate the feature-family totals; the "
         "right panel ranks individual pair features, led by the VHIP-supplied Homology feature."),
        ("Figure 3", "External comparison on identical held-out pairs. Grouped AUROC bars "
         "compare PrecisionPhage, published PHIST, and an in-house RaFAH-inspired proxy "
         "across four regimes; receiver-operating-characteristic curves compare GBM and PHIST."),
        ("Figure 4", "Observed host-taxon coverage versus selected cocktail size. "
         "Model-driven greedy selection is compared with a truth-informed greedy reference "
         "and random selection; the predicted full-cover ILP solution contains 176 phages."),
        ("Figure 5", "Assumption-driven resistance sensitivity simulation over 96 hours. "
         "Total bacterial load and resistant subpopulation are shown for control, monophage, "
         "nonredundant cocktail, and redundant cocktail strategies; both cocktail strategies "
         "rebound to carrying capacity with resistant takeover."),
    ]
    # Set accessibility metadata through python-docx so Word's namespace and
    # compatibility declarations are preserved. Re-serializing document.xml
    # with xml.etree can leave a package that permissive readers accept but
    # Microsoft Word correctly reports as corrupted.
    doc = Document(docx_path)
    properties = [shape._inline.docPr for shape in doc.inline_shapes]
    if len(properties) != len(alt_texts):
        raise AssertionError(
            f"expected {len(alt_texts)} figure descriptions, found {len(properties)}")
    for prop, (title, description) in zip(properties, alt_texts):
        prop.set("title", title)
        prop.set("descr", description)
    doc.save(docx_path)

    with tempfile.TemporaryDirectory() as td:
        staged = Path(td) / docx_path.name
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
                staged, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = (replacements[item.filename].read_bytes()
                        if item.filename in replacements else zin.read(item.filename))
                zout.writestr(item, data)
        shutil.copy2(staged, docx_path)


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    p = doc.paragraphs

    skill = pd.read_csv(RD / "significance_skill.csv")
    cmp = pd.read_csv(RD / "significance_modelcmp.csv")
    leakage = pd.read_csv(RD / "leakage_splits_summary.csv")
    abl_auc = pd.read_csv(RD / "gnn_ablation_auc.csv")
    abl_cmp = pd.read_csv(RD / "gnn_ablation_cmp.csv")
    phist_skill = pd.read_csv(RD / "phist_skill.csv")
    phist_cmp = pd.read_csv(RD / "phist_modelcmp.csv")
    rafah_skill = pd.read_csv(RD / "rafah_skill.csv")
    rafah_cmp = pd.read_csv(RD / "rafah_modelcmp.csv")
    feature_importance = pd.read_csv(RD / "feature_importance.csv").set_index("feature")
    baseline = pd.read_csv(RD / "baseline_summary.csv")
    cold_folds = json.loads(
        (RD / "coldstart_fold_aucs.json").read_text(encoding="utf-8")
    )["combined_unseen"]
    cocktail = json.loads((RD / "cocktail_summary.json").read_text(encoding="utf-8"))
    robustness = {int(x["k"]): x for x in cocktail["robustness"]}
    temporal = json.loads((RD / "temporal_summary.json").read_text(encoding="utf-8"))
    gbm_rows = skill[skill.model == "GBM"].set_index("regime")
    gnn_rows = skill[skill.model == "GNN"].set_index("regime")
    cmp_rows = cmp.set_index("regime")
    leak_rows = leakage[leakage.model == "GBM"].set_index("regime")
    phist_rows = phist_skill[phist_skill.model == "PHIST"].set_index("regime")
    phist_cmp_rows = phist_cmp.set_index("regime")
    rafah_rows = rafah_skill[rafah_skill.model == "RaFAH_style"].set_index("regime")
    rafah_cmp_rows = rafah_cmp.set_index("regime")

    set_text(p[0], "Leakage-controlled evaluation of phage–host interaction prediction with exploratory cocktail design and resistance sensitivity analysis")
    set_text(p[2], "1 West Valley College, Saratoga, CA, United States")
    set_text(p[3], "* Correspondence: Jaivardhan Samadhiya, jsamadhi@mywvm.wvm.edu")
    set_text(p[4], "Article type: Original Research. Manuscript statistics: approximately 4,100 words (main text); 5 figures; 5 tables.")
    set_text(p[8], (
        f"Computational prediction of bacteriophage–host interactions may help prioritize experimental host-range testing, but evaluation can be inflated when related genomes occur in both training and test sets. We present PrecisionPhage, a reproducible computational framework for leakage-controlled interaction prediction and exploratory downstream cocktail design. The direct interaction source contained 8,849 experimentally assayed pairs from three study labels; complete frozen genomes were available for a 1,947-pair NCBI_HR subset (1,488 positive, 459 negative; 1,418 phages and 323 host taxa), which was used for sequence-based modeling. A fixed gradient-boosted tree combined genome composition, source-supplied pair features, and locally computed nucleotide, CRISPR-like, and protein proxies. Pooled AUROC decreased from {gbm_rows.loc['loso_species','auc']:.3f} for eligible unseen host species to {gbm_rows.loc['combined_unseen','auc']:.3f} across five dual cold-start blocks. The tree outperformed an inductive GraphSAGE comparator in the saved row-wise analysis; within the neural architecture, message passing gave its largest gain in dual cold start. PHIST and a RaFAH-inspired proxy were evaluated on identical held-out pairs; these comparisons favor PrecisionPhage but are exploratory because pairs share phage and host entities. Nested-threshold host-cluster-grouped predictions yielded a {cocktail['greedy_min_size']}-phage predicted set cover with {100*robustness[1]['true_cover_>=1']:.1f}% observed at-least-one coverage among {cocktail['n_targets']} eligible host taxa. In a deterministic resistance sensitivity model, both nonredundant and redundant cocktails rebounded and ended with resistant takeover despite an assumption that structurally favors redundancy. These results establish an auditable computational baseline and identify sequence coverage, entity-aware uncertainty, external transfer, and experimental validation as necessary next steps before therapeutic interpretation."
    ))
    set_text(p[12], (
        "Here we prioritize realistic generalization over a single headline accuracy. PrecisionPhage evaluates fixed genome and pair features under taxonomic and sequence-cluster holdouts; fits preprocessing inside training partitions; compares a gradient-boosted tree with GraphSAGE, PHIST, and a RaFAH-inspired proxy on aligned test pairs; and explores how out-of-fold scores behave in set-cover and resistance-sensitivity analyses. The work is Original Research because it reports new computational analyses and negative as well as positive findings. It does not claim wet-lab or clinical validation."
    ))
    set_text(p[15], (
        "We used the 8,849 laboratory-tested virus–host pairs distributed with the Virus-host interactions predictor (VHIP) study (Bastien et al., 2024), comprising the NahantCollection, NCBI_HR, and StaphStudy labels. After canonicalization and de-duplication, the table contained 2,770 positive and 6,079 negative pairs spanning 2,331 phages and 387 host identifiers (380 species labels and 122 genera), with no label conflicts. We did not independently reconstruct the three component datasets. Exact normalized identifiers, plus deterministic aliases for NCBI accession version suffixes, linked entities to frozen FASTAs; fuzzy substring matching was not used. Both genomes resolved for 1,947 pairs (1,488 positive and 459 negative; 1,418 phages and 323 host taxa). All of these pairs came from NCBI_HR. Thus, the full four-feature baseline is multi-study, whereas the sequence, GNN, external-baseline, cocktail, and temporal analyses are single-source."
    ))
    set_text(p[17], (
        "Each genome was summarized by a 212-dimensional node descriptor comprising canonical 4-mer frequencies, codon composition, GC fraction, genome length, and dinucleotide composition. The 24 pair features comprised four locally recomputed composition comparisons (cosine distance, L1 distance, Pearson correlation, and Jaccard similarity); four features supplied by the VHIP table (k3dist, k6dist, GCdiff, and Homology); eight multiscale exact-word proxies at k = 16 and 20; four CRISPR-like spacer-match features; and four six-frame translated-protein 7-mer proxies. The latter 16 features were computed only from each frozen genome pair. Because the four VHIP-supplied features are not regenerated here, the saved 24-feature model is not yet a fully portable predictor for an arbitrary external pair. For the GNN, principal-component analysis and scaling were fitted on the inner-training partition and applied to validation and test data."
    ))
    set_text(p[19], (
        "We used four primary evaluation regimes: leave-one-eligible-host-species-out (LOSO), leave-one-host-sequence-cluster-out, leave-one-phage-sequence-cluster-out, and dual cold start. Taxonomic groups were eligible as test folds only with at least three positive and three negative pairs; 28 of 323 sequence-covered host taxa met this criterion. Sequence groups were generated by an in-house forward-strand bottom-k MinHash sketch (k = 21; 256 hashes), a Mash-form distance transformation, and single-linkage clustering at a 0.05 cutoff (Ondov et al., 2016). This is a custom approximation, not a run of Mash, and the cutoff is not a direct 95% ANI measurement. In each of five dual cold-start blocks, test pairs had both entity clusters assigned to the held-out bin; pairs sharing only one held-out axis were discarded as a leakage buffer. The fixed XGBoost configuration was used for headline results. Nested tuning was evaluated separately but was not selected for the saved headline run. Cocktail probabilities used five-fold stratified group cross-validation grouped by host sequence cluster."
    ))
    set_text(p[21], (
        "The primary model was XGBoost with 400 trees, maximum depth 5, learning rate 0.05, subsample 0.8, and column subsample 0.8 (Chen and Guestrin, 2016). The saved headline GBM probabilities were not post-hoc calibrated. The comparator was a two-layer inductive GraphSAGE encoder with an edge-feature multilayer-perceptron decoder (Hamilton et al., 2017), implemented with scikit-learn, PyTorch, and PyTorch Geometric (Pedregosa et al., 2011; Paszke et al., 2019; Fey and Lenssen, 2019). Within each outer fold, the GNN used a stratified inner split: PCA, scalers, and the message-passing graph were built only from inner-training data and positive edges; early stopping and isotonic calibration used the held-out inner-validation partition (Niculescu-Mizil and Caruana, 2005). An identical neural architecture with an empty graph isolated the contribution of message passing."
    ))
    set_text(p[23], (
        "PHIST was built from its published source and scored covered pairs by exact shared 25-mers (Zielezinski et al., 2022). The published pretrained RaFAH model could not be executed in the original environment. We therefore used a RaFAH-inspired proxy: a fold-internal random forest predicted host genus from feature-hashed amino-acid 6-mer presence after six-frame translation. This proxy does not use RaFAH's released weights, HMM database, or protein clusters and must not be interpreted as a benchmark of the published RaFAH software (Coutinho et al., 2021). Both comparators were evaluated on the same saved test-pair order as the GBM."
    ))
    set_text(p[25], (
        "We report pooled AUROC, AUPRC, expected calibration error (10 bins), and fold-level variation. Saved analyses also include DeLong confidence intervals and paired tests, McNemar tests at a 0.5 threshold, 10,000-row paired bootstrap intervals, and 1,000 row-label permutations with Benjamini–Hochberg correction (DeLong et al., 1988; Benjamini and Hochberg, 1995). These row-wise procedures assume independent observations, which is violated because pairs share phages and hosts. Their p-values and intervals are therefore presented only as exploratory diagnostics; they should not be treated as entity-level confirmatory inference."
    ))
    set_text(p[27], (
        "Five-fold host-sequence-cluster-grouped outer-OOF GBM scores were converted to binary decisions using a separate threshold for each fold. Each threshold maximized F1 on group-aware inner-OOF predictions generated only from that outer fold's training partition; outer test labels did not select their threshold. The predicted matrix and observed-label matrix were kept distinct. Untested cells were unavailable rather than validated negatives. SciPy's mixed-integer linear-programming interface with HiGHS and greedy set cover selected phages on the predicted matrix, and selections were scored against observed labels (Virtanen et al., 2020). Eligible targets were host taxa with at least one observed positive pair. The temporal component was a deterministic sensitivity analysis, not an efficacy validation."
    ))
    set_text(p[29], (
        "Let P be the outer-OOF probability matrix, Âp,h = 1[Pp,h ≥ τf] the predicted binary matrix using the inner-selected threshold for outer fold f, and T the observed-positive matrix. For predicted minimum k-cover, we minimized Σp xp subject to Σp Âp,h xp ≥ min(k, Σp Âp,h) for each eligible host taxon h, with xp ∈ {0,1}. The cap makes the problem feasible when fewer than k predicted phages are available. Greedy selection used marginal predicted gain. All reported coverage values were then calculated with T; cells without an assay contributed no observed coverage. Consequently, a predicted full cover need not achieve complete observed coverage."
    ))
    set_text(p[30], (
        "The sensitivity model tracked susceptible Sh and resistant Rh populations for each host taxon and free phages Vp, following a standard population-dynamics framing (Levin and Bull, 2004). Infection strength was the predicted susceptibility. Parameters were r = 1 h⁻¹, K = 10⁹ CFU/mL, burst size 50, phage decay 0.1 h⁻¹, adsorption 10⁻⁹, resistance cost c = 0.05, mutation rate μ = 10⁻⁷ per division, initial Sh = 10⁶ CFU/mL, and dose 10⁸ PFU/mL per selected phage. Simulations used 193 output points over 96 h (0.5-h spacing). The effective resistance rate was assumed to be μ raised to the number of targeting phages. This independent-resistance assumption structurally favors redundant cocktails and was neither learned from the interaction data nor validated with cross-resistance experiments."
    ))

    software_h = insert_after(p[30], "2.9 Software, reproducibility, and AI assistance", "Heading 2")
    insert_after(software_h, (
        "Analyses used Python 3.11 with locked package versions; release checks record data counts, exact genome resolution, required artifacts, and SHA-256 checksums. OpenAI Codex (GPT-5.6 Sol, OpenAI) assisted with code and reproducibility auditing and with language editing after the initial analyses. The author reviewed and verified all changes, made the scientific decisions, and accepts responsibility for the manuscript and code."
    ), "Normal")

    set_text(p[32], (
        "We first evaluated discrimination as progressively stricter entity holdouts were applied, then examined model class, feature importance, external comparators, predicted set cover, and temporal sensitivity. Results below distinguish the full multi-study four-feature baseline from the 1,947-pair single-source sequence analysis."
    ))
    regime_names = {
        "loso_species": "eligible unseen-species folds",
        "host_cluster": "unseen host clusters",
        "phage_cluster": "unseen phage clusters",
        "combined_unseen": "five dual cold-start blocks",
    }
    headline = ", ".join(
        f"{fmt_ci(gbm_rows.loc[r])} for {regime_names[r]}"
        for r in ("loso_species", "host_cluster", "phage_cluster", "combined_unseen")
    )
    fold_values = list(cold_folds["gbm"].values())
    fold_text = ", ".join(f"{x:.3f}" for x in fold_values)
    cross = baseline[baseline.regime == "cross_study"].set_index("model")
    set_text(p[34], (
        f"For the sequence-covered NCBI_HR subset, pooled GBM AUROC was {headline} (Table 1, Figure 1). Table 2 reports fold means, pooled estimates, and eligible-fold counts. Cold-start fold AUROCs were {fold_text} (mean {sum(fold_values)/len(fold_values):.3f}). Only 28 host taxa met the LOSO class-count criterion, and pooled rows share entities; therefore, the apparently narrow row-wise intervals and chance-test q values overstate the number of independent units. The performance decline and fold spread support the qualitative conclusion that stricter holdouts are harder, while the exact generalization uncertainty remains preliminary. The separate full-table cross-study baseline had mean fold AUROC {cross.loc['GBM','mean_auc']:.3f} for GBM and {cross.loc['EdgeMLP','mean_auc']:.3f} for an edge MLP, indicating poor transfer across the three source labels."
    ))

    pa = abl_auc.set_index(["regime", "model"])
    pc = abl_cmp.set_index("regime")
    qmax = cmp_rows["delong_q_bh"].max()
    set_text(p[36], (
        f"In the freshly rerun saved comparison, the fixed GBM exceeded the GNN in all four regimes; GNN AUROC ranged from {gnn_rows.loc['phage_cluster','auc']:.3f} for unseen phage clusters to {gnn_rows.loc['loso_species','auc']:.3f} across eligible unseen-species folds (Table 1). Row-wise DeLong q values were at most {qmax:.2g}, but shared entities make these exploratory rather than confirmatory. In the architecture ablation, message passing produced positive row-pooled AUROC differences in all four regimes. The largest gain occurred in dual cold start: graph AUROC was {pa.loc[('combined_unseen','GNN_graph'),'auc']:.3f}, compared with {pa.loc[('combined_unseen','GNN_nograph'),'auc']:.3f} without graph edges (difference {pc.loc['combined_unseen','graph_gain']:+.3f}; row-wise q = {pc.loc['combined_unseen','graph_q_bh']:.2g}). The other three graph gains did not survive the exploratory BH threshold. This provides limited evidence that relational information can help the neural comparator under the strictest split, while the fixed GBM remained stronger and the tests are not entity-independent."
    ))
    set_text(p[38], (
        f"Gain-based XGBoost importance assigned {100*feature_importance.loc['Phage genome composition','importance']:.1f}% to phage node composition and {100*feature_importance.loc['Host genome composition','importance']:.1f}% to host node composition; multiscale nucleotide proxies contributed {100*feature_importance.loc['Nucleotide homology','importance']:.1f}%, translated-protein proxies {100*feature_importance.loc['Protein homology','importance']:.1f}%, composition comparisons {100*feature_importance.loc['Composition distance','importance']:.1f}%, and CRISPR-like matches {100*feature_importance.loc['CRISPR matching','importance']:.1f}% (Figure 2). These are model-specific predictive attributions, not evidence that the corresponding biological mechanisms contribute in the same proportions. The low CRISPR-like importance may reflect feature design, genome quality, or data coverage and should not be interpreted as biological irrelevance."
    ))
    set_text(p[40], (
        f"On identical saved test rows, GBM AUROC exceeded PHIST by {phist_cmp_rows.auc_diff.min():.3f}–{phist_cmp_rows.auc_diff.max():.3f} across regimes (Table 3, Figure 3). The RaFAH-inspired proxy was closest for unseen phage clusters ({rafah_rows.loc['phage_cluster','auc']:.3f} versus {rafah_cmp_rows.loc['phage_cluster','auc_gbm']:.3f}) and had AUROC {rafah_rows.loc['combined_unseen','auc']:.3f} in dual cold start. Row-wise DeLong comparisons favor the GBM, but repeated entities invalidate a confirmatory interpretation of the very small q values. PHIST is a genuine external software comparison; the RaFAH-inspired row is only a proxy-method comparison and cannot establish superiority to published RaFAH."
    ))
    set_text(p[42], (
        f"Among {cocktail['n_targets']} observed-positive host taxa, {cocktail['n_predicted_targets']} had at least one predicted covering phage. Predicted set cover selected {cocktail['greedy_min_size']} phages and achieved {100*robustness[1]['true_cover_>=1']:.1f}% observed at-least-one coverage over all {cocktail['n_targets']} taxa. Predicted k = 2 and k = 3 selections contained {robustness[2]['cocktail_size']} and {robustness[3]['cocktail_size']} phages, respectively; their observed at-least-one coverage was {100*robustness[2]['true_cover_>=1']:.1f}% and {100*robustness[3]['true_cover_>=1']:.1f}%, while observed k-fold coverage was {100*robustness[2]['true_cover_>=k']:.1f}% and {100*robustness[3]['true_cover_>=k']:.1f}% (Table 4, Figure 4). Thus, predicted redundancy did not translate into complete observed redundancy. Thresholds were selected without outer-test labels, but the matrix remains sparse and the large solution is not a clinically practical formulation. The result is best viewed as retrospective prioritization behavior on the available assays."
    ))
    set_text(p[43], "3.6 Both cocktail strategies rebound in an assumption-driven sensitivity model")
    outcomes = {x["strategy"]: x for x in temporal["outcomes"]}
    k1, k2 = outcomes["cocktail_k1"], outcomes["robust_k2"]
    set_text(p[44], (
        f"Under the explicit independent-resistance assumption, the k = 1 strategy reached a nadir of {k1['nadir']:.2g} CFU/mL but returned to approximately {k1['end_load']:.1e} CFU/mL with an end resistant fraction of {k1['resistant_frac_end']:.1f}. The k ≥ 2 strategy reached a nadir of {k2['nadir']:.2g} and ended at {k2['end_load']:.1e} CFU/mL with an end resistant fraction of {k2['resistant_frac_end']:.1f}. Both strategies met the predefined rebound criterion and neither prevented modeled resistant takeover (Table 5, Figure 5). The selected five-taxon panel was chosen for high predicted candidate counts, and the μ^n independence rule structurally favors redundancy; failure even under that assumption underscores that this simulation provides no efficacy validation."
    ))
    set_text(p[46], (
        "The central result is a leakage hierarchy rather than a claim of therapeutic readiness. Discrimination decreased when novel phage and host sequence clusters were simultaneously held out, and five cold-start blocks varied substantially. The analysis demonstrates that split design changes the apparent difficulty of phage–host prediction, while also showing that the current sequence-covered evidence is confined to NCBI_HR and a limited set of eligible taxa."
    ))
    set_text(p[47], (
        "Composition-based WIsH, exact-word PHIST, protein-centric RaFAH and CHERRY, and network-integrated predictors capture different signals (Ahlgren et al., 2017; Galiez et al., 2017; Coutinho et al., 2021; Shang and Sun, 2022; Wang et al., 2020; Zielezinski et al., 2022). A recent 27-tool benchmark likewise found strong context dependence across datasets and task formulations (Shang et al., 2025). PrecisionPhage combines several feature families and evaluates them on aligned holdouts. Direct accuracy comparisons to values reported in other papers remain inappropriate unless datasets, candidate-host spaces, negatives, and split units match."
    ))
    set_text(p[48], (
        "The GNN result should be framed narrowly. In this dataset and implementation, message passing yielded modest gains over the otherwise identical no-graph neural model, with its largest gain under dual cold start, but the GNN remained below the fixed GBM in every regime. This pattern may reflect sparse graph structure, model selection, or the dominance of explicit pair features; it neither establishes nor rules out the general value of graph learning. Similarly, the PHIST result is a genuine software benchmark, whereas the RaFAH-inspired proxy identifies a direction for comparison but cannot substitute for running published RaFAH."
    ))
    set_text(p[49], (
        "The downstream analyses are exploratory stress tests of model output. Predicted set cover can rank candidates, but sparse untested cells, thresholds selected from inner-training labels, taxon-level rather than strain-level targets, and very large solutions prevent a therapeutic-cocktail claim. The temporal analysis adds no independent biological observations, and neither cocktail strategy prevented rebound or resistant takeover even though the imposed μ^n rule structurally favors redundancy. A defensible next experiment would freeze a model and threshold before testing all selected phage–strain pairs in a dense independent host-range panel, followed by time-course cross-resistance assays."
    ))
    set_text(p[50], (
        "Limitations include single-source coverage for all 24-feature results; use of four VHIP-supplied features that are not regenerated here; evaluation of only eligible taxonomic groups; five cold-start blocks; repeated phage and host entities that invalidate independent-row inference; no external dense matrix for the headline model; no wet-lab or prospective validation; a RaFAH-inspired proxy rather than published RaFAH; sparse cocktail assays; and an assumption-driven deterministic temporal model. The current work is computational Original Research, but its contribution is methodological and hypothesis-generating rather than evidence of clinical efficacy."
    ))
    set_text(p[52], (
        "PrecisionPhage provides an auditable computational analysis of how taxonomic and sequence-cluster holdouts affect phage–host prediction, together with an architecture ablation and exploratory set-cover and temporal sensitivity studies. The frozen GBM retains discriminatory signal in dual cold start, but sequence coverage, external transfer, entity-aware uncertainty, full published-tool benchmarking, and prospective experimental testing remain necessary before model-selected therapies can be proposed."
    ))

    set_text(p[58], "The author declares that no financial support was received for the research, authorship, or publication of this article.")
    set_text(p[60], (
        "OpenAI Codex (GPT-5.6 Sol, OpenAI) was used for code and reproducibility auditing and for language editing. The author reviewed and verified all changes and accepts full responsibility for the work. No generative AI system is listed as an author."
    ))
    set_text(p[62], (
        "The direct interaction table is available with the VHIP publication (Bastien et al., 2024). Frozen analysis code, configuration, genome staging maps, processed tables, figures, checksums, and reproduction instructions will be available at https://github.com/JaivardhanSamadhiya/frontiers_submission before submission. NCBI accessions are recorded in the repository staging maps."
    ))

    references = [
        "Ahlgren, N. A., Ren, J., Lu, Y. Y., Fuhrman, J. A., and Sun, F. (2017). Alignment-free d2* oligonucleotide frequency dissimilarity measure improves prediction of hosts from metagenomically-derived viral sequences. Nucleic Acids Res. 45, 39–53. doi: 10.1093/nar/gkw1002",
        "Bastien, G. E., Cable, R. N., Batterbee, C., Wing, A. J., Zaman, L., and Duhaime, M. B. (2024). Virus-host interactions predictor (VHIP): machine learning approach to resolve microbial virus-host interaction networks. PLoS Comput. Biol. 20, e1011649. doi: 10.1371/journal.pcbi.1011649",
        "Benjamini, Y., and Hochberg, Y. (1995). Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. Series B 57, 289–300. doi: 10.1111/j.2517-6161.1995.tb02031.x",
        "Chan, B. K., Abedon, S. T., and Loc-Carrillo, C. (2013). Phage cocktails and the future of phage therapy. Future Microbiol. 8, 769–783. doi: 10.2217/fmb.13.47",
        "Chen, T., and Guestrin, C. (2016). XGBoost: a scalable tree boosting system. Proc. 22nd ACM SIGKDD, 785–794. doi: 10.1145/2939672.2939785",
        "Coutinho, F. H., Zaragoza-Solas, A., López-Pérez, M., Barylski, J., Zielezinski, A., Dutilh, B. E., et al. (2021). RaFAH: host prediction for viruses of Bacteria and Archaea based on protein content. Patterns 2, 100274. doi: 10.1016/j.patter.2021.100274",
        "DeLong, E. R., DeLong, D. M., and Clarke-Pearson, D. L. (1988). Comparing the areas under two or more correlated receiver operating characteristic curves: a nonparametric approach. Biometrics 44, 837–845. doi: 10.2307/2531595",
        "Edwards, R. A., McNair, K., Faust, K., Raes, J., and Dutilh, B. E. (2016). Computational approaches to predict bacteriophage–host relationships. FEMS Microbiol. Rev. 40, 258–272. doi: 10.1093/femsre/fuv048",
        "Fey, M., and Lenssen, J. E. (2019). Fast graph representation learning with PyTorch Geometric. ICLR Workshop on Representation Learning on Graphs and Manifolds. arXiv:1903.02428.",
        "Galiez, C., Siebert, M., Enault, F., Vincent, J., and Söding, J. (2017). WIsH: who is the host? Predicting prokaryotic hosts from metagenomic phage contigs. Bioinformatics 33, 3113–3114. doi: 10.1093/bioinformatics/btx383",
        "Hamilton, W. L., Ying, R., and Leskovec, J. (2017). Inductive representation learning on large graphs. Adv. Neural Inf. Process. Syst. 30, 1024–1034.",
        "Levin, B. R., and Bull, J. J. (2004). Population and evolutionary dynamics of phage therapy. Nat. Rev. Microbiol. 2, 166–173. doi: 10.1038/nrmicro822",
        "Murray, C. J. L., Ikuta, K. S., Sharara, F., Swetschinski, L., Robles Aguilar, G., Gray, A., et al. (2022). Global burden of bacterial antimicrobial resistance in 2019: a systematic analysis. Lancet 399, 629–655. doi: 10.1016/S0140-6736(21)02724-0",
        "Niculescu-Mizil, A., and Caruana, R. (2005). Predicting good probabilities with supervised learning. Proc. 22nd Int. Conf. Mach. Learn., 625–632. doi: 10.1145/1102351.1102430",
        "Ondov, B. D., Treangen, T. J., Melsted, P., Mallonee, A. B., Bergman, N. H., Koren, S., et al. (2016). Mash: fast genome and metagenome distance estimation using MinHash. Genome Biol. 17, 132. doi: 10.1186/s13059-016-0997-x",
        "Paszke, A., Gross, S., Massa, F., Lerer, A., Bradbury, J., Chanan, G., et al. (2019). PyTorch: an imperative style, high-performance deep learning library. Adv. Neural Inf. Process. Syst. 32, 8024–8035.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., et al. (2011). Scikit-learn: machine learning in Python. J. Mach. Learn. Res. 12, 2825–2830.",
        "Shang, J., and Sun, Y. (2022). CHERRY: a Computational metHod for accuratE pRediction of virus–pRokarYotic interactions using a graph encoder–decoder model. Brief. Bioinform. 23, bbac182. doi: 10.1093/bib/bbac182",
        "Shang, J., Peng, C., Guan, J., Cai, D., Wang, D., and Sun, Y. (2025). From genomic signals to prediction tools: a critical feature analysis and rigorous benchmark for phage-host prediction. Brief. Bioinform. 26, bbaf626. doi: 10.1093/bib/bbaf626",
        "Torres-Barceló, C., and Hochberg, M. E. (2016). Evolutionary rationale for phages as complements of antibiotics. Trends Microbiol. 24, 249–256. doi: 10.1016/j.tim.2015.12.011",
        "Virtanen, P., Gommers, R., Oliphant, T. E., Haberland, M., Reddy, T., Cournapeau, D., et al. (2020). SciPy 1.0: fundamental algorithms for scientific computing in Python. Nat. Methods 17, 261–272. doi: 10.1038/s41592-019-0686-2",
        "Wang, W., Ren, J., Tang, K., Dart, E., Ignacio-Espinoza, J. C., Fuhrman, J. A., et al. (2020). A network-based integrated framework for predicting virus–prokaryote interactions. NAR Genom. Bioinform. 2, lqaa044. doi: 10.1093/nargab/lqaa044",
        "Zielezinski, A., Deorowicz, S., and Gudyś, A. (2022). PHIST: fast and accurate prediction of prokaryotic hosts from metagenomic viral sequences. Bioinformatics 38, 1447–1449. doi: 10.1093/bioinformatics/btab837",
    ]
    ref_paras = p[64:83]
    for para, text in zip(ref_paras, references[:len(ref_paras)]):
        set_text(para, text)
    cursor = ref_paras[-1]
    for text in references[len(ref_paras):]:
        cursor = insert_after(cursor, text, "Normal")

    # Table 1: one coherent, freshly rerun significance artifact family.
    regimes = ["loso_species", "host_cluster", "phage_cluster", "combined_unseen"]
    labels = ["Eligible unseen species (LOSO)", "Unseen host cluster",
              "Unseen phage cluster", "Both unseen (cold start)"]
    table1 = doc.tables[0]
    for i, (regime, label) in enumerate(zip(regimes, labels), start=1):
        g = skill[(skill.regime == regime) & (skill.model == "GBM")].iloc[0]
        n = skill[(skill.regime == regime) & (skill.model == "GNN")].iloc[0]
        c = cmp[cmp.regime == regime].iloc[0]
        fold_text = (", ".join(f"{x:.3f}" for x in fold_values)
                     if regime == "combined_unseen" else "—")
        vals = [label, fmt_ci(g), fold_text, f"{g.ece:.3f}", fmt_ci(n),
                f"{c.auc_diff:+.3f} ({c.diff_lo:.3f}–{c.diff_hi:.3f})",
                f"{c.delong_q_bh:.2g}*"]
        for j, val in enumerate(vals):
            table1.cell(i, j).text = val
    table1.cell(0, 6).text = "Row-wise DeLong q*"

    # Table 2: fresh leakage hierarchy.
    table2 = doc.tables[1]
    leak_regimes = ["loso_species", "logo_genus", "host_cluster",
                    "phage_cluster", "combined_unseen"]
    leak_labels = ["LOSO (eligible unseen species)", "LOGO (eligible unseen genus)",
                   "Unseen host cluster", "Unseen phage cluster",
                   "Both unseen (cold start)"]
    for i, (regime, label) in enumerate(zip(leak_regimes, leak_labels), start=1):
        row = leak_rows.loc[regime]
        vals = [label, f"{row.mean_auc:.3f} ({row.ci_lo:.3f}–{row.ci_hi:.3f})",
                f"{row.pooled_auc:.3f}", f"{row.ece:.3f}", str(int(row.folds_used))]
        for j, val in enumerate(vals):
            table2.cell(i, j).text = val

    # Table 3: fresh external-comparator values.
    table3 = doc.tables[2]
    table3.cell(0, 3).text = "RaFAH-inspired AUROC"
    table3.cell(0, 5).text = "ΔAUROC vs proxy (row-wise q*)"
    for i, regime in enumerate(("loso_species", "host_cluster", "phage_cluster",
                                "combined_unseen"), start=1):
        ps = phist_rows.loc[regime]
        pc_row = phist_cmp_rows.loc[regime]
        rs = rafah_rows.loc[regime]
        rc_row = rafah_cmp_rows.loc[regime]
        vals = [table3.cell(i, 0).text, f"{pc_row.auc_gbm:.3f}", f"{ps.auc:.3f}",
                f"{rs.auc:.3f}", f"{pc_row.auc_diff:+.3f} ({pc_row.delong_q_bh:.2g})",
                f"{rc_row.auc_diff:+.3f} ({rc_row.delong_q_bh:.2g})"]
        for j, val in enumerate(vals):
            table3.cell(i, j).text = val

    # Fresh cocktail and temporal outcomes in other tables.
    table4 = doc.tables[3]
    table4.cell(0, 2).text = "Host taxa covered ≥1×"
    table4.cell(0, 3).text = "Host taxa covered ≥k×"
    for i, k in enumerate((1, 2, 3), start=1):
        row = robustness[k]
        vals = [str(k), str(row["cocktail_size"]),
                f"{100*row['true_cover_>=1']:.1f}%",
                f"{100*row['true_cover_>=k']:.1f}%"]
        for j, val in enumerate(vals):
            table4.cell(i, j).text = val
    table5 = doc.tables[4]
    row_order = ["control", "monophage", "cocktail_k1", "robust_k2"]
    display = ["No treatment", "Monophage", "Cocktail (k=1)", "Redundant cocktail (k≥2)"]
    for i, (key, label) in enumerate(zip(row_order, display), start=1):
        row = outcomes[key]
        vals = [label, str(row["n_phages"]), f"{row['end_load']:.2e}",
                f"{row['nadir']:.2e}", f"{row['log10_drop']:.2f}",
                f"{row['resistant_frac_end']:.3f}", "Yes" if row["rebound"] else "No"]
        for j, val in enumerate(vals):
            table5.cell(i, j).text = val
    for table in doc.tables:
        mark_header_row(table)

    set_text(p[85], "Table 1. Frozen GBM and freshly rerun GNN performance on aligned test rows. Asterisks mark exploratory row-wise inference; pairs are not independent because they share phage and host entities.")
    set_text(p[87], "Table 2. GBM leakage hierarchy. Taxonomic folds include only groups with at least three positive and three negative pairs; dual cold start contains five blocks.")
    set_text(p[89], "Table 3. External comparisons on aligned test rows. PHIST is the published tool; the RaFAH-inspired proxy is not published RaFAH. Row-wise q values are exploratory.")
    set_text(p[91], "Table 4. Predicted k-redundant selections scored against observed positive labels among eligible host taxa. Untested cells are unavailable.")
    set_text(p[93], "Table 5. Assumption-driven temporal sensitivity outcomes. Both cocktail strategies rebound and end with modeled resistant takeover.")
    set_text(p[97], "Figure 1. GBM and GNN discrimination across taxonomic and sequence-cluster holdouts, with the message-passing ablation. Intervals are row-wise and exploratory because entities repeat.")
    set_text(p[98], "Figure 2. XGBoost gain-based feature importances. These predictive attributions are not proportional biological-mechanism estimates.")
    set_text(p[99], "Figure 3. External comparisons on aligned held-out rows. PHIST is the published software; RaFAH-inspired is an in-house proxy, not released RaFAH.")
    set_text(p[100], "Figure 4. Observed at-least-one coverage as predicted-greedy, truth-oracle, and random selections increase in size among eligible host taxa.")
    set_text(p[101], "Figure 5. Deterministic resistance sensitivity analysis over 96 h. The independent-resistance assumption structurally favors redundant cocktails; both cocktail strategies rebound by the predefined criterion.")

    # Refresh the displayed approximate main-text count.
    all_paras = doc.paragraphs
    intro_i = next(i for i, x in enumerate(all_paras) if x.text == "1 Introduction")
    coi_i = next(i for i, x in enumerate(all_paras) if x.text == "Conflict of Interest")
    main_words = sum(len(re.findall(r"\b[\w–-]+\b", x.text))
                     for x in all_paras[intro_i:coi_i] if not x.style.name.startswith("Heading"))
    set_text(p[4], f"Article type: Original Research. Manuscript statistics: approximately {main_words:,} main-text words; 5 figures; 5 tables.")

    doc.save(OUTPUT)
    replace_embedded_figures(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

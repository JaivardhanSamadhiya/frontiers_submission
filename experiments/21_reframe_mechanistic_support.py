#!/usr/bin/env python3
"""Reframe cocktail and resistance analyses and add a factorial sensitivity table."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_submission_ready.docx"
OUTPUT = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_mechanistic_reframed.docx"
SENSITIVITY = ROOT / "data" / "results_v2" / "temporal_resistance_sensitivity.csv"


def replace_paragraph(doc: Document, starts_with: str, replacement: str) -> None:
    matches = [p for p in doc.paragraphs if p.text.startswith(starts_with)]
    if len(matches) != 1:
        raise AssertionError(f"expected one paragraph starting {starts_with!r}; found {len(matches)}")
    paragraph = matches[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(replacement)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def sci_range(values: pd.Series) -> str:
    return f"{values.min():.2e}–{values.max():.2e}"


def main() -> None:
    if not SOURCE.exists() or not SENSITIVITY.exists():
        raise FileNotFoundError("required manuscript or sensitivity output is missing")
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    sens = pd.read_csv(SENSITIVITY)

    replace_paragraph(
        doc,
        "Leakage-controlled evaluation",
        "Leakage-controlled evaluation of phage–host interaction prediction with exploratory set-cover prioritization and resistance-assumption stress testing",
    )
    replace_paragraph(
        doc,
        "Computational prediction",
        "Computational prediction of bacteriophage–host interactions may help prioritize experimental host-range testing, but evaluation can be inflated when related genomes occur in both training and test sets. We present PrecisionPhage, a reproducible computational framework for leakage-controlled interaction prediction and exploratory downstream set-cover prioritization. The direct interaction source contained 8,849 experimentally assayed pairs from three study labels; complete frozen genomes were available for a 1,947-pair NCBI_HR subset (1,488 positive, 459 negative; 1,418 phages and 323 host taxa), which was used for sequence-based modeling. A fixed gradient-boosted tree combined genome composition, source-supplied pair features, and locally computed nucleotide, CRISPR-like, and protein proxies. Pooled AUROC decreased from 0.959 for eligible unseen host species to 0.785 across five dual cold-start blocks. The tree had higher pooled AUROC than an inductive GraphSAGE comparator in all four evaluated regimes; within the neural architecture, message passing gave its largest descriptive gain in dual cold start. PHIST and a RaFAH-inspired proxy were evaluated on identical held-out pairs and had lower pooled AUROC, but repeated entities preclude confirmatory row-level inference. Nested-threshold host-cluster-grouped predictions yielded a 176-phage predicted set cover with 89.6% observed at-least-one coverage among 230 eligible host taxa. In a 54-run deterministic resistance-assumption stress test, every modeled cocktail trajectory rebounded; resistant takeover occurred in 48 runs, including all complete-cross-resistance settings. These analyses are prioritization and assumption-stress tests, not mechanistic or efficacy validation. The results establish an auditable computational baseline and identify sequence coverage, entity-aware uncertainty, external transfer, and experimental validation as necessary next steps before therapeutic interpretation.",
    )
    replace_paragraph(doc, "2.7 Cocktail optimization", "2.7 Set-cover prioritization and deterministic scenario analysis")
    replace_paragraph(doc, "2.8 Combinatorial optimization", "2.8 Set-cover formulation and deterministic scenario equations")
    replace_paragraph(
        doc,
        "The sensitivity model tracked",
        "The deterministic scenario model tracked susceptible Sₕ and resistant Rₕ populations for each host taxon and free phages Vₚ, following a standard population-dynamics framing (Levin and Bull, 2004). It was not fitted to time-course, adsorption, burst-size, mutation, fitness, or cross-resistance measurements and therefore was not a mechanistic or efficacy-validation model. Infection strength was the predicted susceptibility. The primary parameter setting was r = 1 h⁻¹, K = 10⁹ CFU/mL, burst size 50, phage decay 0.1 h⁻¹, adsorption 10⁻⁹, resistance cost c = 0.05, mutation rate μ = 10⁻⁷ per division, initial Sₕ = 10⁶ CFU/mL, and dose 10⁸ PFU/mL per selected phage. Simulations used 193 output points over 96 h (0.5-h spacing). For n targeting phages, the effective resistance rate was μ^[1 + α(n − 1)], where α = 0 represented complete cross-resistance (μ), α = 0.5 a phenomenological intermediate, and α = 1 independent resistance (μⁿ). A factorial stress test crossed these three α values with μ ∈ {10⁻⁸, 10⁻⁷, 10⁻⁶} and c ∈ {0, 0.05, 0.10} for the k = 1 and k ≥ 2 selections (54 runs). All other parameters were fixed. This grid probes dependence on stated assumptions; it does not estimate biological parameters or validate resistance mechanisms.",
    )
    replace_paragraph(doc, "3.5 Predictions yield", "3.5 Exact optimization of predicted set cover produces large selections")
    replace_paragraph(
        doc,
        "Among 230 observed-positive host taxa",
        "Among 230 observed-positive host taxa, 219 had at least one predicted covering phage. Exact optimization of the thresholded predicted matrix selected 176 phages for k = 1 and achieved 89.6% observed at-least-one coverage over all 230 taxa. Predicted k = 2 and k = 3 selections contained 282 and 358 phages, respectively; their observed at-least-one coverage was 93.0% and 97.0%, while observed k-fold coverage was 45.2% and 31.3% (Table 4, Figure 4). Thus, predicted redundancy did not translate into complete observed redundancy. “Optimal” here refers only to the mathematical set-cover objective on this sparse, thresholded prediction matrix; it does not mean biologically optimal, synergistic, resistance-suppressing, safe, or clinically practical. The result is retrospective prioritization behavior on available assays, not a therapeutic-cocktail result.",
    )
    replace_paragraph(doc, "3.6 Both cocktail strategies", "3.6 Resistance outcomes remain assumption-driven across the sensitivity grid")
    replace_paragraph(
        doc,
        "Under the explicit independent-resistance assumption",
        "At the primary independent-resistance setting (α = 1, μ = 10⁻⁷, c = 0.05), the k = 1 selection reached a nadir of 1.69 × 10⁴ CFU/mL and the k ≥ 2 selection reached 3.98 × 10⁴ CFU/mL; both returned to approximately 1.00 × 10⁹ CFU/mL with end resistant fractions of 1.0 (Table 5, Figure 5). Across the 54-run factorial grid, all 54 trajectories met the predefined rebound criterion and 48 ended with resistant fractions ≥0.5 (Table 6). The k = 1 selection showed takeover in all 27 settings. For k ≥ 2, takeover occurred in 9/9 complete-cross-resistance settings and 6/9 settings under each partial-dependence and independent-resistance assumption; its final loads under the latter two assumptions ranged from 3.99 × 10⁷ to approximately 1.00 × 10⁹ CFU/mL. The selected five-taxon panel was chosen for high predicted candidate counts. These results show sensitivity to the resistance structure but provide no evidence that redundancy prevents rebound or that either selection would be efficacious.",
    )
    replace_paragraph(
        doc,
        "The downstream analyses are exploratory",
        "The downstream analyses are exploratory transformations of model output, not independent mechanistic support. Set-cover optimization can rank candidates, but sparse untested cells, inner-training-selected thresholds, taxon-level rather than strain-level targets, and very large selections prevent a therapeutic-cocktail claim. The deterministic model imposes rather than learns growth, adsorption, burst, mutation, fitness, and cross-resistance parameters. Its original μⁿ rule structurally favored redundancy; relaxing that rule changed some takeover outcomes for k ≥ 2 but did not prevent rebound in any of the 54 settings. Accordingly, neither mathematical redundancy nor the simulated trajectories establish synergy, resistance suppression, or efficacy. A defensible next experiment would freeze the model and threshold before testing every selected phage–strain pair in a dense independent host-range panel, followed by time-course and cross-resistance assays.",
    )
    replace_paragraph(
        doc,
        "Limitations include",
        "Limitations include single-source coverage for all 24-feature results; use of four VHIP-supplied features that are not regenerated here; evaluation of only eligible taxonomic groups; five cold-start blocks; repeated phage and host entities that invalidate independent-row inference; no external dense matrix for the headline model; no wet-lab or prospective validation; a RaFAH-inspired proxy rather than published RaFAH; sparse cocktail assays; and a deterministic scenario model whose growth, adsorption, burst, mutation, fitness, and cross-resistance parameters were not fitted or identified from this dataset. The 54-run grid explores only a bounded set of structural assumptions and cannot establish biological robustness. The current work is computational Original Research, but its contribution is methodological and hypothesis-generating rather than evidence of clinical efficacy.",
    )
    replace_paragraph(
        doc,
        "PrecisionPhage provides",
        "PrecisionPhage provides an auditable computational analysis of how taxonomic and sequence-cluster holdouts affect phage–host prediction, together with an architecture ablation, set-cover prioritization, and deterministic resistance-assumption stress tests. The frozen GBM retains discriminatory signal in dual cold start, but the optimization and simulations do not constitute mechanistic or efficacy evidence. Sequence coverage, external transfer, entity-aware uncertainty, full published-tool benchmarking, and prospective experimental testing remain necessary before model-selected therapies can be proposed.",
    )
    replace_paragraph(
        doc,
        "Table 4.",
        "Table 4. Predicted set-cover selections scored against observed positive labels among eligible host taxa. “Optimal” describes only the mathematical objective on the thresholded predicted matrix; untested cells are unavailable and no therapeutic property is inferred.",
    )
    replace_paragraph(
        doc,
        "Table 5.",
        "Table 5. Primary deterministic scenario outcomes under independent resistance (α = 1), μ = 10⁻⁷, and resistance cost c = 0.05. These trajectories are assumption-driven and are not efficacy validation.",
    )
    replace_paragraph(
        doc,
        "Figure 5.",
        "Figure 5. Primary deterministic resistance scenario over 96 h (α = 1, μ = 10⁻⁷, c = 0.05). The independent-resistance rule structurally favors redundancy, yet both selected cocktails rebound. The curves are an assumption stress test, not mechanistic or efficacy validation.",
    )

    # Add a compact summary of all 54 runs immediately before Figure Legends.
    figure_legends = next(p for p in doc.paragraphs if p.text == "Figure Legends")
    caption = doc.add_paragraph(
        "Table 6. Factorial resistance-assumption stress test. Each row summarizes nine mutation-rate/fitness-cost settings. Takeover denotes an end resistant fraction ≥0.5; every setting rebounded."
    )
    caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=5)
    table.style = doc.tables[-2].style
    headers = ["Set", "Resistance structure (α)", "Rebound\n(n/9)", "Takeover\n(n/9)", "Final-load range (CFU/mL)"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    labels = {
        0.0: "Complete cross-resistance (0)",
        0.5: "Partial dependence (0.5)",
        1.0: "Independent resistance (1)",
    }
    names = {"cocktail_k1": "k = 1", "robust_k2": "k ≥ 2"}
    for (strategy, alpha), group in sens.groupby(["strategy", "resistance_independence"], sort=True):
        row = table.add_row().cells
        values = [
            names[strategy], labels[float(alpha)], f"{int(group['rebound'].sum())}/9",
            f"{int((group['resistant_fraction_end'] >= 0.5).sum())}/9",
            sci_range(group["end_load"]),
        ]
        for cell, value in zip(row, values):
            cell.text = value
    set_repeat_table_header(table.rows[0])
    set_table_widths(table, [0.65, 1.55, 0.90, 0.90, 1.85])
    figure_legends._p.addprevious(caption._p)
    figure_legends._p.addprevious(table._tbl)

    for shape in doc.inline_shapes:
        if shape._inline.docPr.get("title") == "Figure 5":
            shape._inline.docPr.set(
                "descr",
                "Primary deterministic resistance scenario over 96 hours under the independent-resistance assumption. Total bacterial load and resistant subpopulation are shown for control, monophage, k equals 1, and k at least 2 selections. Both selected cocktails rebound; the curves are assumption stress tests, not efficacy validation.",
            )

    # Refresh the Frontiers first-page manuscript statistics after the edits.
    paragraphs = doc.paragraphs
    intro_index = next(i for i, p in enumerate(paragraphs) if p.text == "1 Introduction")
    declarations_index = next(i for i, p in enumerate(paragraphs) if p.text == "Conflict of Interest")
    word_count = sum(
        len(re.findall(r"\b[\w–-]+\b", p.text))
        for p in paragraphs[intro_index:declarations_index]
        if not p.style.name.startswith("Heading")
    )
    replace_paragraph(
        doc,
        "Article type: Original Research.",
        f"Article type: Original Research. Manuscript statistics: approximately {word_count:,} main-text words; 5 figures; 6 tables.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Create the Frontiers submission copy with non-independent inference de-emphasized."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_revised.docx"
OUTPUT = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_submission_ready.docx"


def replace_paragraph(doc: Document, starts_with: str, replacement: str) -> None:
    matches = [p for p in doc.paragraphs if p.text.startswith(starts_with)]
    if len(matches) != 1:
        raise AssertionError(f"expected one paragraph starting {starts_with!r}; found {len(matches)}")
    paragraph = matches[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(replacement)


def remove_column(table, column_index: int) -> None:
    for row in table.rows:
        cell = row.cells[column_index]
        cell._tc.getparent().remove(cell._tc)
    grid = table._tbl.tblGrid
    grid_columns = list(grid.gridCol_lst)
    grid.remove(grid_columns[column_index])


def set_table_widths(table, widths: list[float]) -> None:
    if len(table.columns) != len(widths):
        raise AssertionError("table width specification does not match column count")
    table.autofit = False
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(int(Inches(widths[0]).emu / 635) if False else int(width * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    # The Frontiers template already contains manual section numbers in the
    # heading text. Remove inherited automatic numbering to avoid duplicated
    # labels such as "3.1  2.1" after Word refreshes the document.
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        p_pr = style._element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is not None:
                p_pr.remove(num_pr)
            # A direct numId=0 override suppresses numbering inherited from a
            # linked list style when Word recalculates the document.
            num_pr = OxmlElement("w:numPr")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "0")
            num_pr.append(num_id)
            p_pr.append(num_pr)

    replace_paragraph(
        doc,
        "Computational prediction of bacteriophage–host interactions",
        "Computational prediction of bacteriophage–host interactions may help prioritize experimental host-range testing, but evaluation can be inflated when related genomes occur in both training and test sets. We present PrecisionPhage, a reproducible computational framework for leakage-controlled interaction prediction and exploratory downstream cocktail design. The direct interaction source contained 8,849 experimentally assayed pairs from three study labels; complete frozen genomes were available for a 1,947-pair NCBI_HR subset (1,488 positive, 459 negative; 1,418 phages and 323 host taxa), which was used for sequence-based modeling. A fixed gradient-boosted tree combined genome composition, source-supplied pair features, and locally computed nucleotide, CRISPR-like, and protein proxies. Pooled AUROC decreased from 0.959 for eligible unseen host species to 0.785 across five dual cold-start blocks. The tree had higher pooled AUROC than an inductive GraphSAGE comparator in all four evaluated regimes; within the neural architecture, message passing gave its largest descriptive gain in dual cold start. PHIST and a RaFAH-inspired proxy were evaluated on identical held-out pairs and had lower pooled AUROC, but repeated entities preclude confirmatory row-level inference. Nested-threshold host-cluster-grouped predictions yielded a 176-phage predicted set cover with 89.6% observed at-least-one coverage among 230 eligible host taxa. In a deterministic resistance sensitivity model, both nonredundant and redundant cocktails rebounded and ended with resistant takeover despite an assumption that structurally favors redundancy. These results establish an auditable computational baseline and identify sequence coverage, entity-aware uncertainty, external transfer, and experimental validation as necessary next steps before therapeutic interpretation.",
    )
    replace_paragraph(
        doc,
        "We report pooled AUROC, AUPRC",
        "We report pooled AUROC, AUPRC, expected calibration error (10 bins), and variation across held-out folds or blocks. The repository retains legacy exploratory calculations using DeLong intervals and paired tests, McNemar tests at a 0.5 threshold, 10,000-row paired bootstrap intervals, and 1,000 row-label permutations with Benjamini–Hochberg correction (DeLong et al., 1988; Benjamini and Hochberg, 1995). These procedures treat rows as independent or exchangeable, an assumption violated because pairs share phages and hosts. We therefore omit their p-values and q-values from the main results and do not use them for confirmatory claims. Pooled AUROC differences are descriptive effect sizes; fold-level estimates are emphasized as the available assessment of between-group variation. The five dual cold-start blocks are too few to support precise entity-level inference.",
    )
    replace_paragraph(
        doc,
        "For the sequence-covered NCBI_HR subset",
        "For the sequence-covered NCBI_HR subset, pooled GBM AUROC was 0.959 for eligible unseen-species folds, 0.950 for unseen host clusters, 0.847 for unseen phage clusters, and 0.785 for five dual cold-start blocks (Table 1, Figure 1). Table 2 reports fold means, fold-level intervals, pooled estimates, and eligible-fold counts. Cold-start fold AUROCs were 0.894, 0.688, 0.867, 0.765, and 0.785 (mean 0.800). Only 28 host taxa met the LOSO class-count criterion, and just five independent blocks contributed to dual cold start. The performance decline and fold spread support the qualitative conclusion that stricter holdouts are harder, while the exact generalization uncertainty remains preliminary. The current locked-environment full-table cross-study rerun had mean fold AUROC 0.486 for GBM and 0.608 for an edge MLP, indicating limited transfer across the three source labels.",
    )
    replace_paragraph(
        doc,
        "3.2 A gradient-boosted tree outperforms",
        "3.2 The gradient-boosted tree has higher pooled AUROC than the graph neural network",
    )
    replace_paragraph(
        doc,
        "In the freshly rerun saved comparison",
        "In the freshly rerun saved comparison, the fixed GBM had higher pooled AUROC than the GNN in all four regimes; descriptive differences ranged from +0.090 to +0.113, and GNN AUROC ranged from 0.752 for unseen phage clusters to 0.869 across eligible unseen-species folds (Table 1). Because rows share phage and host entities, these differences are not presented as confirmatory hypothesis tests. In the architecture ablation, message passing produced positive row-pooled AUROC differences in all four regimes. The largest descriptive gain occurred in dual cold start: graph AUROC was 0.672, compared with 0.576 without graph edges (difference +0.097). This pattern suggests that relational information may help the neural comparator under the strictest split, but it does not establish an entity-independent effect; the fixed GBM retained the higher pooled AUROC.",
    )
    replace_paragraph(
        doc,
        "3.4 PrecisionPhage exceeds external baselines",
        "3.4 PrecisionPhage has higher pooled AUROC than the evaluated baselines",
    )
    replace_paragraph(
        doc,
        "On identical saved test rows, GBM AUROC exceeded PHIST",
        "On identical saved test rows, GBM pooled AUROC was higher than PHIST by 0.104–0.275 across regimes (Table 3, Figure 3). The RaFAH-inspired proxy was closest for unseen phage clusters (0.745 versus 0.847) and had AUROC 0.414 in dual cold start. These are descriptive comparisons: repeated phage and host entities prevent confirmatory row-level inference. PHIST is a genuine external software comparison; the RaFAH-inspired row is only a proxy-method comparison and cannot establish superiority to published RaFAH.",
    )
    replace_paragraph(
        doc,
        "The direct interaction table is available with the VHIP publication",
        "The direct interaction table is available with the VHIP publication (Bastien et al., 2024). Frozen analysis code, configuration, genome staging maps, processed tables, figures, checksums, and reproduction instructions are available at https://github.com/JaivardhanSamadhiya/frontiers_submission. Large genome FASTA files are excluded because of their size; their NCBI accessions and staging identifiers are recorded in the repository.",
    )
    replace_paragraph(
        doc,
        "Table 1. Frozen GBM",
        "Table 1. Frozen GBM and freshly rerun GNN performance on aligned test rows. AUROCs and differences are descriptive pooled estimates; pairs are not independent because they share phage and host entities. Fold-level variation is reported in Table 2.",
    )
    replace_paragraph(
        doc,
        "Table 3. External comparisons",
        "Table 3. Descriptive external comparisons on aligned test rows. PHIST is the published tool; the RaFAH-inspired proxy is not published RaFAH. Differences are pooled effect sizes, not confirmatory tests, because phage and host entities repeat.",
    )
    replace_paragraph(
        doc,
        "Figure 1. GBM and GNN discrimination",
        "Figure 1. GBM and GNN discrimination across taxonomic and sequence-cluster holdouts, with the message-passing ablation. Pooled intervals shown in the figure are exploratory row-wise summaries because entities repeat; fold-level variation in Table 2 is emphasized for interpretation.",
    )

    table1 = doc.tables[0]
    table1.cell(0, 1).text = "GBM pooled AUROC"
    table1.cell(0, 4).text = "GNN pooled AUROC"
    table1.cell(0, 5).text = "Pooled ΔAUROC"
    table1.cell(0, 6).text = "Status"
    for row in table1.rows[1:]:
        row.cells[1].text = row.cells[1].text.split(" (")[0]
        row.cells[4].text = row.cells[4].text.split(" (")[0]
        row.cells[5].text = row.cells[5].text.split(" (")[0]
        row.cells[6].text = "Descriptive"
    set_table_widths(table1, [1.15, 0.85, 1.25, 0.55, 0.85, 0.90, 0.95])

    table3 = doc.tables[2]
    table3.cell(0, 4).text = "Descriptive ΔAUROC vs PHIST"
    table3.cell(0, 5).text = "Descriptive ΔAUROC vs proxy"
    for row in table3.rows[1:]:
        row.cells[4].text = row.cells[4].text.split(" (")[0]
        row.cells[5].text = row.cells[5].text.split(" (")[0]
    set_table_widths(table3, [1.25, 1.00, 0.85, 1.00, 1.20, 1.20])

    # Keep each table caption with its table. The template's final "Figures"
    # heading is redundant after the complete Figure Legends section and its
    # anchored layout obscures most of the word in Word's PDF export, so remove
    # it and begin the figure appendix cleanly with Figure 1.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True
    figures_heading = next(p for p in doc.paragraphs if p.text == "Figures")
    figures_heading._element.getparent().remove(figures_heading._element)
    # The complete captions and figure numbers are already provided in the
    # Figure Legends section. Remove redundant standalone labels before the
    # embedded images; Word's template can partially obscure a label when a
    # figure is pushed to the top of a new page.
    for paragraph in list(doc.paragraphs):
        if re.fullmatch(r"Figure [1-5]", paragraph.text):
            paragraph._element.getparent().remove(paragraph._element)

    paragraphs = doc.paragraphs
    intro_index = next(i for i, p in enumerate(paragraphs) if p.text == "1 Introduction")
    declarations_index = next(i for i, p in enumerate(paragraphs) if p.text == "Conflict of Interest")
    word_count = sum(
        len(re.findall(r"\b[\w–-]+\b", p.text))
        for p in paragraphs[intro_index:declarations_index]
        if not p.style.name.startswith("Heading")
    )
    replace_paragraph(
        doc,
        "Article type: Original Research.",
        f"Article type: Original Research. Manuscript statistics: approximately {word_count:,} main-text words; 5 figures; 5 tables.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

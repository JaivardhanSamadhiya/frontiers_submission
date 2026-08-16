#!/usr/bin/env python3
"""Integrate the locked NCBI_HR -> StaphStudy validation into the manuscript."""
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
SOURCE = (ROOT / "submission" /
          "PrecisionPhage_Frontiers_Original_Research_unfitted_parameter_reframed.docx")
OUTPUT = (ROOT / "submission" /
          "PrecisionPhage_Frontiers_Original_Research_frozen_external_validated.docx")
RESULTS = ROOT / "data" / "results_v2" / "external_staph_validation.json"


def replace_paragraph(doc: Document, starts_with: str, replacement: str) -> None:
    matches = [p for p in doc.paragraphs if p.text.startswith(starts_with)]
    if len(matches) != 1:
        raise AssertionError(
            f"expected one paragraph starting {starts_with!r}; found {len(matches)}"
        )
    paragraph = matches[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(replacement)


def insert_after(paragraph, text: str, style: str | None = None):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new._p)
    return new


def clear_direct_numbering(paragraph) -> None:
    """Disable list numbering inherited from the template's Heading 2 style."""
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        num_id = OxmlElement("w:numId")
        num_pr.append(num_id)
    num_id.set(qn("w:val"), "0")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    dxa = [int(round(width * 1440)) for width in widths]
    total = sum(dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        for cell, width in zip(row.cells, dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True


def add_reference_before(doc: Document, before_starts: str, reference: str) -> None:
    target = next(p for p in doc.paragraphs if p.text.startswith(before_starts))
    paragraph = doc.add_paragraph(reference)
    paragraph.style = target.style
    target._p.addprevious(paragraph._p)


def main() -> None:
    if not SOURCE.is_file() or not RESULTS.is_file():
        raise FileNotFoundError("required manuscript or validation results are missing")
    result = json.loads(RESULTS.read_text(encoding="utf-8"))
    full = result["models"]["full"]
    sequence = result["models"]["sequence_only"]
    four = result["models"]["vhip_four_feature"]

    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    replace_paragraph(
        doc,
        "Leakage-controlled",
        "Leakage-controlled phage–host interaction prediction reveals source-dependent transfer",
    )
    replace_paragraph(
        doc,
        "Computational prediction",
        "Computational prediction of bacteriophage–host interactions may prioritize experimental host-range testing, but related genomes and source structure can inflate evaluation. We present PrecisionPhage, a reproducible leakage-controlled framework using 8,849 experimentally assayed pairs from three study labels. Within the 1,947-pair sequence-covered NCBI_HR subset (1,488 positive and 459 negative), pooled fixed-tree AUROC decreased from 0.959 for eligible unseen host species to 0.785 across five dual cold-start blocks. We then locked the feature pipeline and model before retraining-free evaluation on all 1,053 sequence-covered StaphStudy pairs (333 positive and 720 negative). The full representation achieved AUROC 0.637 (two-way sequence-cluster bootstrap 95% CI 0.498–0.739) and AUPRC 0.391 (0.201–0.591) against 31.6% positive prevalence. A sequence-only sensitivity model excluding four source-supplied features was similar (AUROC 0.639), whereas those four features alone gave AUROC 0.444. The full model remained poorly calibrated (expected calibration error 0.499). On 805 pairs beyond the prespecified phage and host similarity threshold, AUROC was 0.644. Nested-threshold predictions within NCBI_HR yielded a 176-phage retrospective set cover with 89.6% observed at-least-one coverage. An unfitted deterministic parameter grid was retained only as an assumption-propagation illustration. These results support a within-source leakage hierarchy and modest but uncertain cross-source ranking, not a generally calibrated predictor, resistance mechanism, or therapeutic efficacy claim.",
    )
    replace_paragraph(
        doc,
        "Here we separate questions",
        "Here we separate questions requiring different evidence. First, a leave-one-study-out audit tests transfer across three source labels using the four features present for every pair. Second, within sequence-covered NCBI_HR, taxonomic and sequence-cluster holdouts compare a fixed gradient-boosted tree with GraphSAGE, PHIST, and a RaFAH-inspired proxy. Third, a protocol locked before prediction evaluates the unchanged full sequence pipeline retraining-free on sequence-covered StaphStudy pairs. Out-of-fold NCBI_HR scores are additionally examined in retrospective set cover. A deterministic resistance grid is only an ancillary illustration of imposed assumptions. This computational Original Research reports positive and negative validation results but does not claim wet-lab, prospective, or clinical validation.",
    )
    replace_paragraph(
        doc,
        "We used the 8,849 laboratory-tested",
        "We used the 8,849 laboratory-tested virus–host pairs distributed with the Virus-host interactions predictor (VHIP) study (Bastien et al., 2024), comprising the NahantCollection, NCBI_HR, and StaphStudy labels. After canonicalization and de-duplication, the table contained 2,770 positive and 6,079 negative pairs spanning 2,331 phages and 387 host identifiers, with no label conflicts. Exact normalized identifiers, plus deterministic accession aliases, linked entities to FASTAs; fuzzy substring matching was prohibited. The original frozen bundle resolved both genomes for 1,947 NCBI_HR pairs (1,488 positive and 459 negative; 1,418 phages and 323 host taxa). For the source-held-out test, the published VHIP example bundle linked all 1,053 StaphStudy pairs to 39 phage and 27 host reference genomes. StaphStudy derives from the experimentally assayed wastewater-phage study of Göller et al. (2021). The test source was excluded from all fitting of the sequence model.",
    )
    replace_paragraph(
        doc,
        "Each genome was summarized",
        "Each genome was summarized by a 212-dimensional node descriptor comprising canonical 4-mer frequencies, codon composition, GC fraction, genome length, and dinucleotide composition. The 24 pair features comprised four locally recomputed composition comparisons; four VHIP-supplied features (k3dist, k6dist, GCdiff, and Homology); eight exact-word proxies at k = 16 and 20; four CRISPR-like spacer-match features; and four six-frame translated-protein 7-mer proxies. The latter 16 depended only on each genome pair. The VHIP-supplied StaphStudy table was byte-identical to the frozen interaction file and matched all four columns exactly by phage–host key. A prespecified sensitivity model omitted those four columns to test whether transfer depended on them. For the GNN, principal-component analysis and scaling were fitted on inner-training data and applied unchanged to validation and test data.",
    )

    # New Methods subsection, inserted before the existing Models subsection.
    models_heading = next(p for p in doc.paragraphs if p.text == "2.4 Models")
    validation_heading = doc.add_paragraph(
        "2.4 Frozen cross-source sequence validation", style=models_heading.style
    )
    clear_direct_numbering(validation_heading)
    validation_text = doc.add_paragraph(
        "Before generating held-out predictions, we locked the test cohort, feature schema, fixed XGBoost configuration, exclusion rules, and metrics in a dated protocol. A StandardScaler and the fixed 400-tree model were fitted once on all 1,947 NCBI_HR rows without access to StaphStudy labels, then applied without retraining or threshold optimization to all 1,053 StaphStudy rows. Primary outcomes were pooled AUROC, AUPRC, Brier score, and 10-bin expected calibration error. Ninety-five-percent intervals used 2,000 valid two-way bootstrap replicates that independently resampled phage and host sequence clusters; row-wise resampling was not used. Secondary analyses excluded the 10 test host species identifiers seen in training (663 pairs) and excluded any pair for which either genome fell at or below the prespecified 0.05 Mash-form distance to a training genome (805 pairs). Exact-sequence overlap was zero on both axes; four test phages and four test hosts fell at or below the distance threshold. Because StaphStudy was already a component of VHRnet and its labels had appeared in the earlier four-feature audit, this evaluation is source-held-out but not prospectively blinded or independently collected.",
        style=doc.styles["Normal"],
    )
    models_heading._p.addprevious(validation_heading._p)
    models_heading._p.addprevious(validation_text._p)
    for old, new in (
        ("2.4 Models", "2.5 Models"),
        ("2.5 External baselines", "2.6 External baselines"),
        ("2.6 Statistical analysis", "2.7 Statistical analysis"),
        ("2.7 Set-cover prioritization", "2.8 Set-cover prioritization and ancillary deterministic illustration"),
        ("2.8 Set-cover formulation", "2.9 Set-cover formulation and illustrative scenario equations"),
        ("2.9 Software", "2.10 Software, reproducibility, and AI assistance"),
    ):
        replace_paragraph(doc, old, new)
    replace_paragraph(
        doc,
        "We report pooled AUROC",
        "We report pooled AUROC, AUPRC, expected calibration error (10 bins), and variation across held-out folds or blocks. For the frozen StaphStudy evaluation, uncertainty was quantified by two-way sequence-cluster bootstrap over both entity axes. The repository retains legacy exploratory DeLong, McNemar, row-bootstrap, and row-permutation calculations, but these assume independent or exchangeable rows even though pairs share phages and hosts (DeLong et al., 1988; Benjamini and Hochberg, 1995). Their p-values and q-values are omitted from the main results and not used for confirmatory claims. Within-source pooled AUROC differences remain descriptive; fold-level variation is emphasized, and five dual cold-start blocks are too few for precise entity-level inference.",
    )
    replace_paragraph(
        doc,
        "Analyses used Python 3.11",
        "Analyses used Python 3.11 with locked package versions; release checks record data counts, exact genome resolution, required artifacts, and SHA-256 checksums. The dated cross-source protocol, external predictions, model artifacts, overlap distances, and two-way bootstrap summaries are retained with the analysis. OpenAI Codex (GPT-5.6 Sol, OpenAI) assisted with code and reproducibility auditing and with language editing after the initial analyses. The author reviewed and verified all changes, made the scientific decisions, and accepts responsibility for the manuscript and code.",
    )
    replace_paragraph(
        doc,
        "We first evaluated transfer",
        "We first evaluated source transfer with the full-table four-feature audit, then evaluated stricter entity holdouts within sequence-covered NCBI_HR. The newly locked test applied the full sequence representation from NCBI_HR to StaphStudy without retraining. Model-class, comparator, set-cover, and deterministic-scenario analyses otherwise remain based on NCBI_HR and are interpreted accordingly.",
    )

    # New Results subsection after the existing 3.1 result paragraph.
    result_31 = next(p for p in doc.paragraphs if p.text.startswith("In the full-table leave-one-source-out audit"))
    result_heading = insert_after(
        result_31,
        "3.2 Frozen source-held-out evaluation shows modest ranking and poor calibration",
        style="Heading 2",
    )
    clear_direct_numbering(result_heading)
    insert_after(
        result_heading,
        "The frozen fixed GBM trained on 1,947 NCBI_HR pairs and applied once to 1,053 StaphStudy pairs achieved AUROC 0.637 (two-way sequence-cluster bootstrap 95% CI 0.498–0.739) and AUPRC 0.391 (0.201–0.591), compared with 31.6% positive prevalence (Table 8). The sequence-only sensitivity model excluding k3dist, k6dist, GCdiff, and Homology was similar (AUROC 0.639; AUPRC 0.392), whereas those four source-supplied features alone transferred poorly (AUROC 0.444; AUPRC 0.270). The full model was not transportably calibrated (Brier score 0.466; ECE 0.499). All 39 test phage identifiers and all exact genome sequences were absent from training. After excluding 10 reused host species identifiers, full-model AUROC was 0.625 on 663 pairs. After excluding pairs with either entity at or below the prespecified sequence-distance threshold, AUROC was 0.644 and AUPRC 0.427 on 805 pairs (35.2% positive). Across eligible groups, median per-phage and per-host AUROCs were 0.625 and 0.663, respectively, but ranges were wide. Thus, the full sequence representation carried modest ranking information across this source boundary, while uncertainty and calibration preclude deployment or broad generalization claims.",
        style="Normal",
    )
    for old, new in (
        ("3.2 Within NCBI_HR", "3.3 Within NCBI_HR, the gradient-boosted tree has higher pooled AUROC than the graph neural network"),
        ("3.3 Genome composition", "3.4 Genome composition and nucleotide homology dominate"),
        ("3.4 Within NCBI_HR", "3.5 Within NCBI_HR, PrecisionPhage has higher pooled AUROC than the evaluated baselines"),
        ("3.5 Exact optimization", "3.6 Exact optimization of predicted set cover produces large selections"),
        ("3.6 Prespecified equations", "3.7 Prespecified equations yield conditional rebound across the illustrative grid"),
    ):
        replace_paragraph(doc, old, new)

    replace_paragraph(
        doc,
        "The central result is conditional",
        "The central result now has two parts. Within sequence-covered NCBI_HR, discrimination decreased as taxonomic and sequence-cluster separation became stricter. Across the NCBI_HR-to-StaphStudy boundary, the frozen full representation retained modest discrimination (AUROC 0.637), but its cluster-aware interval included 0.5 and calibration deteriorated sharply. The close sequence-only result shows that the four VHIP-supplied columns did not create the observed ranking signal. Together, these analyses show that leakage control changes apparent within-source difficulty and that some ranking information transfers to this held-out source, while the magnitude and probability scale remain source-dependent.",
    )
    replace_paragraph(
        doc,
        "Composition-based WIsH",
        "Composition-based WIsH, exact-word PHIST, protein-centric RaFAH and CHERRY, and network-integrated predictors capture different signals (Ahlgren et al., 2017; Galiez et al., 2017; Coutinho et al., 2021; Shang and Sun, 2022; Wang et al., 2020; Zielezinski et al., 2022). A 27-tool benchmark likewise found strong context dependence across datasets and task formulations (Shang et al., 2025). The present results reinforce that warning: source prevalence ranged widely, the four-feature transfer audit was weak, and the full sequence model's ranking exceeded chance only imprecisely while its probabilities were poorly calibrated. A model may therefore preserve ordering across a source boundary without preserving clinically usable probabilities or thresholds.",
    )
    replace_paragraph(
        doc,
        "The GNN result should be framed narrowly",
        "The GNN result remains within-source. Message passing modestly improved the otherwise identical no-graph neural model within NCBI_HR, but the GNN remained below the fixed GBM in every regime and was not frozen for the new StaphStudy test. The new external result therefore validates only the fixed tree pipeline, not the model-class ranking, PHIST comparison, RaFAH-inspired proxy, or downstream set-cover selections. Those comparisons should not be generalized beyond their aligned NCBI_HR test rows.",
    )
    replace_paragraph(
        doc,
        "The downstream analyses are transformations",
        "The downstream analyses are transformations of NCBI_HR model output, not independent evidence of generalization or mechanism. The external ranking result does not retroactively validate the 176-phage set cover, its thresholds, or the deterministic resistance equations. Sparse untested cells, taxon-level targets, very large selections, and poor probability transport prevent a therapeutic-cocktail claim. The equations impose rather than learn resistance and kinetic parameters; their universal rebound within the sampled grid cannot estimate rebound probability, validate a mechanism, compare therapies, or strengthen the introductory AMR rationale. Biological claims would require measured adsorption, killing, mutation, fitness, cross-resistance, and prospective time courses.",
    )
    replace_paragraph(
        doc,
        "The principal limitation is that all 24-feature",
        "The frozen StaphStudy test addresses the previous absence of cross-source sequence evaluation, but it does not eliminate the generalization limitation. StaphStudy was already part of VHRnet and the earlier four-feature audit, so the test was source-held-out rather than prospectively blinded or independently collected. It contains 39 phages, 27 host reference genomes, and only 15 phage sequence clusters; 10 host species identifiers recurred in training, and four test genomes on each axis fell within the prespecified similarity threshold, although the strict 805-pair subset gave a similar result. The full-model AUROC interval included chance and calibration was poor. Further limitations include evaluation of only eligible internal groups, five cold-start blocks, repeated entities, no wet-lab or prospective validation, a RaFAH-inspired proxy rather than published RaFAH, sparse cocktail assays, and unfitted resistance parameters. The work supports a leakage and transportability analysis, not a deployable predictor, resistance model, or efficacy claim.",
    )
    replace_paragraph(
        doc,
        "PrecisionPhage provides",
        "PrecisionPhage provides an auditable within-NCBI_HR leakage hierarchy and a frozen, retraining-free NCBI_HR-to-StaphStudy sequence evaluation. The full representation showed modest but uncertain cross-source ranking and poor probability calibration; it did not validate downstream set cover, resistance biology, or therapeutic efficacy. Independent prospective sequence-covered matrices and empirical experiments remain necessary before predictive deployment or biological claims.",
    )
    replace_paragraph(
        doc,
        "The direct interaction table is available",
        "The direct interaction table is available with the VHIP publication (Bastien et al., 2024), and the StaphStudy genomes are available through the VHIP example bundle and GenBank accessions MZ417315–MZ417354 reported by Göller et al. (2021). Frozen analysis code, the dated external-validation protocol, processed features, predictions, model artifacts, overlap audits, tables, figures, checksums, and reproduction instructions are available at https://github.com/JaivardhanSamadhiya/frontiers_submission. Large genome FASTA files are excluded because of size; accession and staging identifiers are recorded in the repository.",
    )

    add_reference_before(
        doc,
        "Hamilton, W. L.",
        "Göller, P. C., Elsener, T., Lorgé, D., Radulovic, N., Bernardi, V., Naumann, A., et al. (2021). Multi-species host range of staphylococcal phages isolated from wastewater. Nat. Commun. 12, 6965. doi: 10.1038/s41467-021-27037-6",
    )

    # Add compact Table 8 before the figure-legend block.
    figure_legends = next(p for p in doc.paragraphs if p.text == "Figure Legends")
    caption = doc.add_paragraph(
        "Table 8. Frozen NCBI_HR-to-StaphStudy evaluation. Confidence intervals are two-way sequence-cluster bootstrap intervals (2,000 valid replicates). Sequence-only excludes the four VHIP-supplied columns; the four-feature row uses only those columns."
    )
    caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=7)
    table.style = doc.tables[0].style
    headers = ["Model", "Pairs", "Positive", "AUROC (95% CI)",
               "AUPRC (95% CI)", "Brier", "ECE"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = [
        ("Full", full),
        ("Sequence-only", sequence),
        ("Four supplied", four),
    ]
    for label, item in rows:
        pooled = item["pooled"]
        ci = item["ci_95"]
        values = [
            label,
            f"{pooled['n']:,}",
            f"{100 * pooled['prevalence']:.1f}%",
            f"{pooled['roc_auc']:.3f} ({ci['roc_auc']['lo']:.3f}–{ci['roc_auc']['hi']:.3f})",
            f"{pooled['pr_auc']:.3f} ({ci['pr_auc']['lo']:.3f}–{ci['pr_auc']['hi']:.3f})",
            f"{pooled['brier']:.3f}",
            f"{pooled['ece']:.3f}",
        ]
        for cell, value in zip(table.add_row().cells, values):
            cell.text = value
    set_repeat_header(table.rows[0])
    set_table_geometry(table, [1.20, 0.55, 0.65, 1.30, 1.30, 0.75, 0.75])
    figure_legends._p.addprevious(caption._p)
    figure_legends._p.addprevious(table._tbl)

    # Update first-page statistics after all edits.
    paragraphs = doc.paragraphs
    intro_index = next(i for i, p in enumerate(paragraphs) if p.text == "1 Introduction")
    declarations_index = next(i for i, p in enumerate(paragraphs)
                              if p.text == "Conflict of Interest")
    word_count = sum(
        len(re.findall(r"\b[\w–-]+\b", p.text))
        for p in paragraphs[intro_index:declarations_index]
        if not p.style.name.startswith("Heading")
    )
    replace_paragraph(
        doc,
        "Article type: Original Research.",
        f"Article type: Original Research. Manuscript statistics: approximately {word_count:,} main-text words; 5 figures; 8 tables.",
    )

    # Keep appendix table captions with their tables; prevent a caption-only page.
    for paragraph in doc.paragraphs:
        if re.match(r"Table [1-8]\.", paragraph.text):
            paragraph.paragraph_format.keep_with_next = True

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

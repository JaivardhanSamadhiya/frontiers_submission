#!/usr/bin/env python3
"""Make the NCBI_HR scope primary and add the held-out-source audit."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_mechanistic_reframed.docx"
OUTPUT = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_single_source_reframed.docx"
DETAIL = ROOT / "data" / "results_v2" / "cross_study_detail.csv"


def replace_paragraph(doc: Document, starts_with: str, replacement: str) -> None:
    matches = [p for p in doc.paragraphs if p.text.startswith(starts_with)]
    if len(matches) != 1:
        raise AssertionError(f"expected one paragraph starting {starts_with!r}; found {len(matches)}")
    paragraph = matches[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(replacement)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def main() -> None:
    if not SOURCE.is_file() or not DETAIL.is_file():
        raise FileNotFoundError("required manuscript or cross-source detail is missing")
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    detail = pd.read_csv(DETAIL)

    replace_paragraph(
        doc,
        "Leakage-controlled",
        "Leakage-controlled phage–host interaction prediction within a single data source reveals limited cross-source transfer",
    )
    replace_paragraph(
        doc,
        "Computational prediction",
        "Computational prediction of bacteriophage–host interactions may help prioritize experimental host-range testing, but evaluation can be inflated by related genomes and source-specific structure. We present PrecisionPhage, a reproducible computational framework for leakage-controlled interaction prediction. The input contained 8,849 experimentally assayed pairs from three study labels, but complete frozen genomes were available only for a 1,947-pair NCBI_HR subset (1,488 positive, 459 negative; 1,418 phages and 323 host taxa). A full-table leave-one-source-out audit using the four common source-supplied features showed limited transfer: mean fold AUROC was 0.486 for a gradient-boosted tree and 0.608 for an edge multilayer perceptron, and held-out-source class prevalence ranged from 6.1% to 81.9%. All 24-feature sequence analyses were therefore interpreted strictly within NCBI_HR. Within that source, pooled gradient-boosted-tree AUROC decreased from 0.959 for eligible unseen host species to 0.785 across five dual cold-start blocks. The tree had higher pooled AUROC than an inductive GraphSAGE comparator in all four NCBI_HR regimes; PHIST and a RaFAH-inspired proxy were also evaluated only on identical NCBI_HR held-out pairs. Nested-threshold predictions yielded a 176-phage retrospective set cover with 89.6% observed at-least-one coverage among 230 eligible NCBI_HR host taxa. In a 54-run deterministic resistance-assumption stress test derived from those predictions, every trajectory rebounded and 48 ended with resistant takeover. Thus, the study supports a within-source leakage hierarchy and an auditable negative cross-source result; it does not establish cross-database generalization, mechanism, or therapeutic efficacy.",
    )
    replace_paragraph(
        doc,
        "Here we prioritize realistic generalization",
        "Here we separate two questions that require different evidence. First, a leave-one-study-out audit tests transfer across the three source labels using only the four features available for every pair. Second, within the sequence-covered NCBI_HR subset, PrecisionPhage evaluates fixed genome and pair features under taxonomic and sequence-cluster holdouts; compares a gradient-boosted tree with GraphSAGE, PHIST, and a RaFAH-inspired proxy on aligned test pairs; and examines out-of-fold scores in set-cover and resistance-assumption stress tests. The resulting leakage hierarchy is explicitly a within-NCBI_HR finding. The work is computational Original Research because it reports new reproducible analyses and negative as well as positive findings; it does not claim cross-database, wet-lab, or clinical validation.",
    )
    replace_paragraph(
        doc,
        "We used four primary evaluation regimes",
        "For the full 8,849-pair table, the cross-source audit held out each of NCBI_HR, NahantCollection, and StaphStudy in turn, trained on the other two labels, and used only k3dist, k6dist, GCdiff, and Homology, the four source-supplied features present for every pair. We evaluated the fixed XGBoost model and an edge-feature multilayer perceptron, reporting each held-out source separately because prevalence and score calibration differed sharply across sources. This audit does not test the 24-feature sequence model. Within the 1,947-pair NCBI_HR subset, we used four evaluation regimes: leave-one-eligible-host-species-out (LOSO), leave-one-host-sequence-cluster-out, leave-one-phage-sequence-cluster-out, and dual cold start. Taxonomic groups were eligible as test folds only with at least three positive and three negative pairs; 28 of 323 sequence-covered host taxa met this criterion. Sequence groups were generated by an in-house forward-strand bottom-k MinHash sketch (k = 21; 256 hashes), a Mash-form distance transformation, and single-linkage clustering at a 0.05 cutoff (Ondov et al., 2016). This is a custom approximation, not a run of Mash, and the cutoff is not a direct 95% ANI measurement. In each of five dual cold-start blocks, test pairs had both entity clusters assigned to the held-out bin; pairs sharing only one held-out axis were discarded as a leakage buffer. The fixed XGBoost configuration was used for headline within-source results. Nested tuning was evaluated separately but was not selected for the saved headline run. Set-cover probabilities used five-fold stratified group cross-validation grouped by host sequence cluster, also within NCBI_HR.",
    )
    replace_paragraph(
        doc,
        "We first evaluated discrimination",
        "We first evaluated transfer across source labels on the full four-feature table, then evaluated progressively stricter entity holdouts within the sequence-covered NCBI_HR subset. Model-class, feature-importance, external-comparator, set-cover, and deterministic-scenario results all use that single-source subset and are interpreted accordingly.",
    )
    replace_paragraph(
        doc,
        "3.1 Performance declines",
        "3.1 Cross-source transfer is limited, and within-source performance declines under stricter holdouts",
    )
    replace_paragraph(
        doc,
        "For the sequence-covered NCBI_HR subset",
        "In the full-table leave-one-source-out audit, held-out-source positive prevalence ranged from 6.1% for NahantCollection to 81.9% for NCBI_HR (Table 1). GBM AUROC was 0.472 with NCBI_HR held out, 0.419 with NahantCollection held out, and 0.568 with StaphStudy held out (mean 0.486). Edge-MLP AUROC was 0.693, 0.560, and 0.570, respectively (mean 0.608). These results show source heterogeneity and do not validate transfer of the 24-feature sequence model. Within the sequence-covered NCBI_HR subset, pooled GBM AUROC was 0.959 for eligible unseen-species folds, 0.950 for unseen host clusters, 0.847 for unseen phage clusters, and 0.785 for five dual cold-start blocks (Table 2, Figure 1). Table 3 reports fold means, fold-level intervals, pooled estimates, and eligible-fold counts. Cold-start fold AUROCs were 0.894, 0.688, 0.867, 0.765, and 0.785 (mean 0.800). Only 28 host taxa met the LOSO class-count criterion, and just five independent blocks contributed to dual cold start. The decline and fold spread support the qualitative conclusion that stricter holdouts are harder within NCBI_HR; they do not establish that the same hierarchy or performance transfers to another database.",
    )
    replace_paragraph(
        doc,
        "3.2 The gradient-boosted tree",
        "3.2 Within NCBI_HR, the gradient-boosted tree has higher pooled AUROC than the graph neural network",
    )
    replace_paragraph(
        doc,
        "In the freshly rerun saved comparison",
        "Within NCBI_HR, the fixed GBM had higher pooled AUROC than the GNN in all four regimes; descriptive differences ranged from +0.090 to +0.113, and GNN AUROC ranged from 0.752 for unseen phage clusters to 0.869 across eligible unseen-species folds (Table 2). Because rows share phage and host entities, these differences are not presented as confirmatory hypothesis tests. In the architecture ablation, message passing produced positive row-pooled AUROC differences in all four NCBI_HR regimes. The largest descriptive gain occurred in dual cold start: graph AUROC was 0.672, compared with 0.576 without graph edges (difference +0.097). This pattern concerns model behavior within this source only; it neither establishes cross-source transfer nor the general value of graph learning.",
    )
    replace_paragraph(
        doc,
        "3.4 PrecisionPhage has higher",
        "3.4 Within NCBI_HR, PrecisionPhage has higher pooled AUROC than the evaluated baselines",
    )
    replace_paragraph(
        doc,
        "On identical saved test rows",
        "On identical saved NCBI_HR test rows, GBM pooled AUROC was higher than PHIST by 0.104–0.275 across regimes (Table 4, Figure 3). The RaFAH-inspired proxy was closest for unseen phage clusters (0.745 versus 0.847) and had AUROC 0.414 in dual cold start. These are within-source descriptive comparisons: repeated phage and host entities prevent confirmatory row-level inference, and the design provides no evidence of cross-source superiority. PHIST is a genuine external software comparison; the RaFAH-inspired row is only a proxy-method comparison and cannot establish superiority to published RaFAH.",
    )
    replace_paragraph(
        doc,
        "Among 230 observed-positive host taxa",
        "Within NCBI_HR, 219 of 230 observed-positive host taxa had at least one predicted covering phage. Exact optimization of the thresholded predicted matrix selected 176 phages for k = 1 and achieved 89.6% observed at-least-one coverage over all 230 taxa. Predicted k = 2 and k = 3 selections contained 282 and 358 phages, respectively; their observed at-least-one coverage was 93.0% and 97.0%, while observed k-fold coverage was 45.2% and 31.3% (Table 5, Figure 4). Thus, predicted redundancy did not translate into complete observed redundancy. “Optimal” here refers only to the mathematical set-cover objective on this sparse, thresholded, single-source prediction matrix; it does not mean biologically optimal, transferable, synergistic, resistance-suppressing, safe, or clinically practical. The result is retrospective prioritization behavior on available NCBI_HR assays, not a therapeutic-cocktail result.",
    )
    replace_paragraph(
        doc,
        "At the primary independent-resistance setting",
        "All susceptibility inputs to the deterministic scenario came from the single-source NCBI_HR analysis. At the primary independent-resistance setting (α = 1, μ = 10⁻⁷, c = 0.05), the k = 1 selection reached a nadir of 1.69 × 10⁴ CFU/mL and the k ≥ 2 selection reached 3.98 × 10⁴ CFU/mL; both returned to approximately 1.00 × 10⁹ CFU/mL with end resistant fractions of 1.0 (Table 6, Figure 5). Across the 54-run factorial grid, all 54 trajectories met the predefined rebound criterion and 48 ended with resistant fractions ≥0.5 (Table 7). The k = 1 selection showed takeover in all 27 settings. For k ≥ 2, takeover occurred in 9/9 complete-cross-resistance settings and 6/9 settings under each partial-dependence and independent-resistance assumption. These calculations cannot demonstrate performance outside NCBI_HR and provide no evidence that redundancy prevents rebound or that either selection would be efficacious.",
    )
    replace_paragraph(
        doc,
        "The central result is a leakage hierarchy",
        "The central result is conditional: within the sequence-covered NCBI_HR subset, discrimination decreased as taxonomic and sequence-cluster separation became stricter. The full-table four-feature audit is the only direct cross-source test, and its weak, source-dependent AUROCs prevent extrapolation of the 24-feature leakage hierarchy, model rankings, or downstream selections to other databases. The analysis therefore demonstrates that split design changes apparent difficulty within one source, while identifying cross-source transfer as unresolved rather than achieved.",
    )
    replace_paragraph(
        doc,
        "Composition-based WIsH",
        "Composition-based WIsH, exact-word PHIST, protein-centric RaFAH and CHERRY, and network-integrated predictors capture different signals (Ahlgren et al., 2017; Galiez et al., 2017; Coutinho et al., 2021; Shang and Sun, 2022; Wang et al., 2020; Zielezinski et al., 2022). A recent 27-tool benchmark likewise found strong context dependence across datasets and task formulations (Shang et al., 2025). The present source audit reinforces that warning: class prevalence, candidate spaces, assay design, and feature distributions differ across NCBI_HR, NahantCollection, and StaphStudy. Consequently, aligned holdouts within NCBI_HR improve internal validity but cannot substitute for retraining-free evaluation on a separately assembled, sequence-covered external matrix.",
    )
    replace_paragraph(
        doc,
        "The GNN result should be framed narrowly",
        "The GNN result should be framed narrowly. Within NCBI_HR and this implementation, message passing yielded modest gains over the otherwise identical no-graph neural model, but the GNN remained below the fixed GBM in every regime. Because neither model was tested with the full 24-feature representation outside NCBI_HR, this ranking may be source-specific. Similarly, the PHIST result is a genuine within-source software benchmark, whereas the RaFAH-inspired proxy identifies a direction for comparison but cannot substitute for running published RaFAH on an external sequence-covered panel.",
    )
    replace_paragraph(
        doc,
        "The downstream analyses are exploratory",
        "The downstream analyses are exploratory transformations of single-source model output, not independent evidence of generalization or mechanism. Set-cover optimization can rank NCBI_HR candidates, but sparse untested cells, inner-training-selected thresholds, taxon-level rather than strain-level targets, very large selections, and unknown transfer to other sources prevent a therapeutic-cocktail claim. The deterministic model imposes rather than learns biological parameters, and none of its 54 settings prevented rebound. A defensible next validation would freeze the complete feature pipeline, model, and threshold before retraining-free testing on a separately assembled, sequence-covered, dense phage–strain matrix from another source; prospective time-course and cross-resistance assays would then address biology rather than database transfer alone.",
    )
    replace_paragraph(
        doc,
        "Limitations include",
        "The principal limitation is that all 24-feature, GBM/GNN, external-comparator, set-cover, and deterministic-scenario results come from 1,947 NCBI_HR pairs. The multi-study audit uses only four source-supplied features and performs poorly enough that it cannot bridge this evidence gap. Additional limitations include inability to regenerate those four features here; evaluation of only eligible taxonomic groups; five cold-start blocks; repeated entities that invalidate independent-row inference; absence of an external dense sequence-covered matrix; no wet-lab or prospective validation; a RaFAH-inspired proxy rather than published RaFAH; sparse cocktail assays; and unfitted deterministic resistance parameters. The work is computational Original Research, but its contribution is a within-source methodological analysis and a documented transfer limitation, not a generally validated predictor or evidence of clinical efficacy.",
    )
    replace_paragraph(
        doc,
        "PrecisionPhage provides",
        "PrecisionPhage provides an auditable within-NCBI_HR analysis of how taxonomic and sequence-cluster holdouts affect phage–host prediction, together with a negative full-table cross-source audit. The results support the methodological conclusion that leakage control changes apparent performance within this source; they do not establish cross-database generalization of the 24-feature model or its downstream set-cover and resistance-scenario outputs. A frozen, retraining-free evaluation on a separately assembled sequence-covered interaction matrix is the essential next step before broader predictive or therapeutic claims.",
    )

    # Renumber the existing tables so the new cross-source table is first cited.
    caption_replacements = [
        ("Table 6.", "Table 7. Factorial resistance-assumption stress test derived from NCBI_HR predictions. Each row summarizes nine mutation-rate/fitness-cost settings. Takeover denotes an end resistant fraction ≥0.5; every setting rebounded."),
        ("Table 5.", "Table 6. Primary deterministic scenario outcomes derived from NCBI_HR predictions under independent resistance (α = 1), μ = 10⁻⁷, and resistance cost c = 0.05. These trajectories are assumption-driven and are not efficacy validation."),
        ("Table 4.", "Table 5. Predicted set-cover selections scored against observed positive labels within NCBI_HR. “Optimal” describes only the mathematical objective on the sparse thresholded matrix; untested cells are unavailable and no therapeutic or cross-source property is inferred."),
        ("Table 3.", "Table 4. Descriptive external comparisons on aligned NCBI_HR test rows. PHIST is the published tool; the RaFAH-inspired proxy is not published RaFAH. Differences are pooled effect sizes, not confirmatory tests, and do not establish cross-source superiority."),
        ("Table 2.", "Table 3. GBM leakage hierarchy within NCBI_HR. Taxonomic folds include only groups with at least three positive and three negative pairs; dual cold start contains five blocks."),
        ("Table 1.", "Table 2. Frozen GBM and freshly rerun GNN performance on aligned NCBI_HR test rows. AUROCs and differences are descriptive pooled estimates; pairs are not independent because they share phage and host entities. Fold-level variation is reported in Table 3."),
    ]
    for old, new in caption_replacements:
        replace_paragraph(doc, old, new)
    replace_paragraph(
        doc,
        "Figure 1.",
        "Figure 1. GBM and GNN discrimination across taxonomic and sequence-cluster holdouts within NCBI_HR, with the message-passing ablation. Pooled intervals shown in the figure are exploratory row-wise summaries because entities repeat; fold-level variation in Table 3 is emphasized for interpretation.",
    )

    # Insert Table 1 immediately before the pre-existing table appendix.
    old_table_2_caption = next(p for p in doc.paragraphs if p.text.startswith("Table 2."))
    old_table_2_caption.paragraph_format.page_break_before = True
    table_6_caption = next(p for p in doc.paragraphs if p.text.startswith("Table 6."))
    table_6_caption.paragraph_format.page_break_before = True
    caption = doc.add_paragraph(
        "Table 1. Leave-one-source-out audit on all 8,849 pairs using only the four source-supplied features available across NCBI_HR, NahantCollection, and StaphStudy. This is not a cross-source evaluation of the 24-feature sequence model."
    )
    caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=7)
    table.style = doc.tables[0].style
    headers = ["Held-out source", "Model", "Pairs", "Positive", "ROC AUC", "PR AUC", "ECE"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    source_order = {"NCBI_HR": 0, "NahantCollection": 1, "StaphStudy": 2}
    model_order = {"GBM": 0, "EdgeMLP": 1}
    detail = detail.sort_values(
        ["held_out_source", "model"],
        key=lambda s: s.map(source_order if s.name == "held_out_source" else model_order),
    )
    for _, record in detail.iterrows():
        row = table.add_row().cells
        source_label = {
            "NCBI_HR": "NCBI_HR",
            "NahantCollection": "Nahant",
            "StaphStudy": "Staph",
        }[record["held_out_source"]]
        model_label = "Edge MLP" if record["model"] == "EdgeMLP" else record["model"]
        values = [
            source_label, model_label, f"{int(record['n_pairs']):,}",
            f"{100 * record['positive_fraction']:.1f}%", f"{record['auroc']:.3f}",
            f"{record['auprc']:.3f}", f"{record['ece']:.3f}",
        ]
        for cell, value in zip(row, values):
            cell.text = value
    set_repeat_header(table.rows[0])
    set_table_widths(table, [1.45, 0.80, 0.60, 0.75, 0.70, 0.70, 0.60])
    old_table_2_caption._p.addprevious(caption._p)
    old_table_2_caption._p.addprevious(table._tbl)

    # Update the first-page count after all edits.
    paragraphs = doc.paragraphs
    intro_index = next(i for i, p in enumerate(paragraphs) if p.text == "1 Introduction")
    declarations_index = next(i for i, p in enumerate(paragraphs) if p.text == "Conflict of Interest")
    word_count = sum(
        len(re.findall(r"\b[\w–-]+\b", p.text))
        for p in paragraphs[intro_index:declarations_index]
        if not p.style.name.startswith("Heading")
    )
    replace_paragraph(
        doc,
        "Article type: Original Research.",
        f"Article type: Original Research. Manuscript statistics: approximately {word_count:,} main-text words; 5 figures; 7 tables.",
    )
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Demote the unfitted resistance grid from biological evidence to illustration."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_single_source_reframed.docx"
OUTPUT = ROOT / "submission" / "PrecisionPhage_Frontiers_Original_Research_unfitted_parameter_reframed.docx"


def replace_paragraph(doc: Document, starts_with: str, replacement: str) -> None:
    matches = [p for p in doc.paragraphs if p.text.startswith(starts_with)]
    if len(matches) != 1:
        raise AssertionError(f"expected one paragraph starting {starts_with!r}; found {len(matches)}")
    paragraph = matches[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(replacement)


def main() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    replace_paragraph(
        doc,
        "Computational prediction",
        "Computational prediction of bacteriophage–host interactions may help prioritize experimental host-range testing, but evaluation can be inflated by related genomes and source-specific structure. We present PrecisionPhage, a reproducible computational framework for leakage-controlled interaction prediction. The input contained 8,849 experimentally assayed pairs from three study labels, but complete frozen genomes were available only for a 1,947-pair NCBI_HR subset (1,488 positive, 459 negative; 1,418 phages and 323 host taxa). A full-table leave-one-source-out audit using the four common source-supplied features showed limited transfer: mean fold AUROC was 0.486 for a gradient-boosted tree and 0.608 for an edge multilayer perceptron, and held-out-source class prevalence ranged from 6.1% to 81.9%. All 24-feature sequence analyses were therefore interpreted strictly within NCBI_HR. Within that source, pooled gradient-boosted-tree AUROC decreased from 0.959 for eligible unseen host species to 0.785 across five dual cold-start blocks. The tree had higher pooled AUROC than an inductive GraphSAGE comparator in all four NCBI_HR regimes; PHIST and a RaFAH-inspired proxy were also evaluated only on identical NCBI_HR held-out pairs. Nested-threshold predictions yielded a 176-phage retrospective set cover with 89.6% observed at-least-one coverage among 230 eligible NCBI_HR host taxa. An ancillary deterministic parameter-grid illustration was not fitted to resistance data and was not used as biological or therapeutic evidence. Thus, the study supports a within-source leakage hierarchy and an auditable negative cross-source result; it does not establish cross-database generalization, resistance biology, or therapeutic efficacy.",
    )
    replace_paragraph(
        doc,
        "The global rise of antimicrobial resistance",
        "The global rise of antimicrobial resistance, associated with an estimated 1.27 million deaths in 2019, has renewed interest in bacteriophages as precision antibacterials (Murray et al., 2022). Because most phages infect a narrow range of hosts, therapeutic research requires matching phages to the infecting strain and empirically evaluating combinations intended to broaden coverage or limit resistance (Chan et al., 2013; Torres-Barceló and Hochberg, 2016). Computational interaction scores may prioritize laboratory testing, but they do not by themselves demonstrate resistance suppression or therapeutic benefit. Experimentally screening phage–host pairs remains slow, motivating careful evaluation of genome-based prioritization methods (Edwards et al., 2016).",
    )
    replace_paragraph(
        doc,
        "Here we separate two questions",
        "Here we separate questions that require different evidence. First, a leave-one-study-out audit tests transfer across the three source labels using only the four features available for every pair. Second, within the sequence-covered NCBI_HR subset, PrecisionPhage evaluates fixed genome and pair features under taxonomic and sequence-cluster holdouts and compares a gradient-boosted tree with GraphSAGE, PHIST, and a RaFAH-inspired proxy. Out-of-fold scores are additionally examined in retrospective set cover. A deterministic resistance parameter grid is included only as an ancillary illustration of how imposed assumptions propagate through equations; it is excluded from claims about validation, mechanism, AMR, or efficacy. The work is computational Original Research because it reports new reproducible analyses and negative as well as positive findings; it does not claim cross-database, wet-lab, or clinical validation.",
    )
    replace_paragraph(
        doc,
        "2.7 Set-cover prioritization",
        "2.7 Set-cover prioritization and ancillary deterministic illustration",
    )
    replace_paragraph(
        doc,
        "Five-fold host-sequence-cluster-grouped outer-OOF",
        "Five-fold host-sequence-cluster-grouped outer-OOF GBM scores were converted to binary decisions using a separate threshold for each fold. Each threshold maximized F1 on group-aware inner-OOF predictions generated only from that outer fold's training partition; outer test labels did not select their threshold. The predicted matrix and observed-label matrix were kept distinct. Untested cells were unavailable rather than validated negatives. SciPy's mixed-integer linear-programming interface with HiGHS and greedy set cover selected phages on the predicted matrix, and selections were scored against observed labels (Virtanen et al., 2020). Eligible targets were host taxa with at least one observed positive pair. The deterministic temporal component was an ancillary mathematical illustration and was not included in evidence for predictive validation, resistance biology, or therapeutic efficacy.",
    )
    replace_paragraph(
        doc,
        "2.8 Set-cover formulation",
        "2.8 Set-cover formulation and illustrative scenario equations",
    )
    replace_paragraph(
        doc,
        "The deterministic scenario model tracked",
        "The deterministic scenario equations tracked susceptible Sₕ and resistant Rₕ populations for each host taxon and free phages Vₚ, following a standard population-dynamics framing (Levin and Bull, 2004). No time-course, adsorption, burst-size, mutation, fitness-cost, or cross-resistance measurements were available for fitting, calibration, or validation. Infection strength was assigned from predicted susceptibility rather than measured killing. The primary illustrative setting was r = 1 h⁻¹, K = 10⁹ CFU/mL, burst size 50, phage decay 0.1 h⁻¹, adsorption 10⁻⁹, resistance cost c = 0.05, mutation rate μ = 10⁻⁷ per division, initial Sₕ = 10⁶ CFU/mL, and dose 10⁸ PFU/mL per selected phage. Simulations used 193 output points over 96 h. For n targeting phages, the imposed effective resistance rate was μ^[1 + α(n − 1)], where α = 0 represented complete cross-resistance, α = 0.5 a phenomenological intermediate, and α = 1 independent resistance. A prespecified grid crossed these α values with μ ∈ {10⁻⁸, 10⁻⁷, 10⁻⁶} and c ∈ {0, 0.05, 0.10} for the k = 1 and k ≥ 2 selections (54 runs), while all other parameters remained fixed. The grid samples assumptions; it neither estimates a plausible parameter distribution nor supports statistical or biological inference. Any common outcome across the grid is therefore a conditional property of these equations, parameter bounds, initial conditions, and rebound definition—not a validated prediction about phage resistance.",
    )
    replace_paragraph(
        doc,
        "We first evaluated transfer",
        "We first evaluated transfer across source labels on the full four-feature table, then evaluated progressively stricter entity holdouts within the sequence-covered NCBI_HR subset. Model-class, feature-importance, external-comparator, and set-cover results all use that single-source subset. The deterministic parameter grid is reported separately as an ancillary illustration and is not treated as validation evidence.",
    )
    replace_paragraph(
        doc,
        "3.6 Resistance outcomes",
        "3.6 Prespecified equations yield conditional rebound across the illustrative grid",
    )
    replace_paragraph(
        doc,
        "All susceptibility inputs to the deterministic scenario",
        "This section describes numerical behavior under the prespecified equations and grid; it does not test a biological resistance hypothesis. All susceptibility inputs came from the single-source NCBI_HR analysis. At the primary setting (α = 1, μ = 10⁻⁷, c = 0.05), the k = 1 and k ≥ 2 selections reached nadirs of 1.69 × 10⁴ and 3.98 × 10⁴ CFU/mL, respectively, before returning to approximately 1.00 × 10⁹ CFU/mL with end resistant fractions of 1.0 (Table 6, Figure 5). All 54 grid trajectories met the chosen rebound criterion, and 48 ended with resistant fractions ≥0.5 (Table 7). These counts summarize only the sampled parameter combinations. Because mutation, fitness cost, cross-resistance, infection strength, and other kinetic quantities were neither fitted nor assigned empirical distributions, the frequency 54/54 is not an estimated probability of rebound and must not be interpreted as evidence that resistance or treatment failure would occur in vitro, in vivo, or clinically.",
    )
    replace_paragraph(
        doc,
        "The downstream analyses are exploratory",
        "The downstream analyses are transformations of single-source model output, not independent evidence of generalization or mechanism. Set-cover optimization can rank NCBI_HR candidates, but sparse untested cells, inner-training-selected thresholds, taxon-level targets, very large selections, and unknown transfer prevent a therapeutic-cocktail claim. The deterministic equations impose rather than learn resistance and kinetic parameters. Their universal rebound within the chosen grid cannot estimate the likelihood of rebound, validate a resistance mechanism, compare therapeutic strategies, or strengthen the introductory AMR rationale; it only documents behavior under those assumptions. External validation should first freeze the prediction pipeline for retraining-free testing on a separately assembled sequence-covered matrix. Any biological resistance claim would additionally require measured adsorption, killing, mutation, fitness, and cross-resistance data and prospective time courses.",
    )
    replace_paragraph(
        doc,
        "The principal limitation is that all 24-feature",
        "The principal limitation is that all 24-feature, GBM/GNN, external-comparator, set-cover, and scenario inputs come from 1,947 NCBI_HR pairs. The multi-study audit uses only four source-supplied features and cannot bridge this evidence gap. Additional limitations include inability to regenerate those four features here; evaluation of only eligible taxonomic groups; five cold-start blocks; repeated entities that invalidate independent-row inference; absence of an external dense sequence-covered matrix; no wet-lab or prospective validation; a RaFAH-inspired proxy rather than published RaFAH; and sparse cocktail assays. Most importantly for the scenario section, the resistance and kinetic parameters are unfitted, their grid bounds are not an empirical uncertainty distribution, and the equations omit biological processes that could alter dynamics. The work is computational Original Research, but its contribution is a within-source methodological analysis and documented transfer limitation—not a validated predictor, resistance model, or efficacy study.",
    )
    replace_paragraph(
        doc,
        "PrecisionPhage provides",
        "PrecisionPhage provides an auditable within-NCBI_HR analysis of how taxonomic and sequence-cluster holdouts affect phage–host prediction, together with a negative full-table cross-source audit. The results support the methodological conclusion that leakage control changes apparent performance within this source; they do not establish cross-database generalization of the 24-feature model. The set-cover analysis is retrospective prioritization, and the unfitted resistance grid is solely an assumption-propagation illustration. Neither provides biological or therapeutic validation. A frozen, retraining-free external evaluation is the essential next predictive step, while resistance or efficacy claims require empirical parameterization and prospective experiments.",
    )
    replace_paragraph(
        doc,
        "Table 6.",
        "Table 6. Primary illustrative trajectory derived from NCBI_HR predictions under imposed independent resistance (α = 1), μ = 10⁻⁷, and resistance cost c = 0.05. Parameters were not fitted; these outputs are conditional equation behavior, not biological prediction or efficacy validation.",
    )
    replace_paragraph(
        doc,
        "Table 7.",
        "Table 7. Conditional outputs of the prespecified resistance-parameter grid derived from NCBI_HR predictions. Each row summarizes nine imposed mutation-rate/fitness-cost settings. The 9/9 counts are not probability estimates because parameters were neither fitted nor sampled from empirical distributions.",
    )
    replace_paragraph(
        doc,
        "Figure 5.",
        "Figure 5. Primary illustrative deterministic trajectory over 96 h under imposed parameters (α = 1, μ = 10⁻⁷, c = 0.05). The curves show conditional equation behavior and are not fitted resistance predictions, mechanistic evidence, or efficacy validation.",
    )

    for shape in doc.inline_shapes:
        if shape._inline.docPr.get("title") == "Figure 5":
            shape._inline.docPr.set(
                "descr",
                "Illustrative deterministic trajectories for total bacterial load and resistant subpopulation under imposed, unfitted parameters. Curves are shown for control, monophage, k equals 1, and k at least 2 selections. The figure documents conditional equation behavior and is not a biological resistance prediction or efficacy validation.",
            )

    # Remove empty spacer paragraphs between the completed legend block and
    # embedded figure paragraphs. With line numbering enabled, a lone spacer
    # can otherwise be pushed onto an otherwise blank page.
    figure_5_legend = next(p for p in doc.paragraphs if p.text.startswith("Figure 5."))
    after_legend = False
    for paragraph in list(doc.paragraphs):
        if paragraph._p is figure_5_legend._p:
            after_legend = True
            continue
        if (after_legend and not paragraph.text.strip()
                and not paragraph._p.xpath(".//w:drawing")
                and not paragraph._p.xpath(".//w:pict")):
            paragraph._element.getparent().remove(paragraph._element)

    paragraphs = doc.paragraphs
    intro_index = next(i for i, p in enumerate(paragraphs) if p.text == "1 Introduction")
    declarations_index = next(i for i, p in enumerate(paragraphs) if p.text == "Conflict of Interest")
    word_count = sum(
        len(re.findall(r"\b[\w–-]+\b", p.text))
        for p in paragraphs[intro_index:declarations_index]
        if not p.style.name.startswith("Heading")
    )
    replace_paragraph(
        doc,
        "Article type: Original Research.",
        f"Article type: Original Research. Manuscript statistics: approximately {word_count:,} main-text words; 5 figures; 7 tables.",
    )
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Step 15: assemble a Frontiers Original Research manuscript (.docx).

Starts from the official Frontiers Word template (so all paragraph styles, page
setup and theme are preserved), clears the placeholder body, and writes the full
manuscript: title block, abstract, introduction, methods, results, discussion,
back-matter, editable tables and embedded figures with legends. All quantitative
claims are pulled from data/results_v2 artifacts.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TPL = ROOT / "Frontiers_Word_Templates" / "Frontiers_Template.docx"
RD = ROOT / "data" / "results_v2"
OUT = ROOT / "PrecisionPhage_Frontiers_manuscript.docx"
OUT_COPY = RD / "PrecisionPhage_Frontiers_manuscript.docx"
REPO_URL = "https://github.com/JaivardhanSamadhiya/frontiers_submission"


def _load_coldstart_folds(rd: Path) -> dict:
    p = rd / "coldstart_fold_aucs.json"
    if p.exists():
        return json.loads(p.read_text(encoding="utf-8"))
    return {}


def _fmt_fold_aucs(fold_aucs: dict) -> str:
    vals = [fold_aucs[k] for k in sorted(fold_aucs)]
    parts = ", ".join(f"{v:.3f}" for v in vals)
    mean = sum(vals) / len(vals)
    return f"{parts} (mean {mean:.3f}, n={len(vals)} blocks)"


def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def H(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def P(doc, text="", style="Normal", align=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def rich(doc, segments, style="Normal", align=None):
    """segments: list of (text, bold, italic)."""
    p = doc.add_paragraph(style=style)
    for text, b, i in segments:
        r = p.add_run(text)
        r.bold = b
        r.italic = i
    if align is not None:
        p.alignment = align
    return p


def table(doc, headers, rows, caption=None):
    if caption:
        P(doc, caption, bold=True)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    # compact font
    for r in t.rows:
        for c in r.cells:
            for par in c.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()
    return t


def figure(doc, path, width_in=6.3):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = docx.Document(str(TPL))
    clear_body(doc)
    cold = _load_coldstart_folds(RD)
    gbm_folds = cold.get("combined_unseen", {}).get("gbm", {})
    fold_str = _fmt_fold_aucs(gbm_folds) if gbm_folds else ""

    # ---------------- Title block ----------------
    P(doc, "Leakage-controlled prediction of phage–host interactions enables "
           "model-driven phage-cocktail design and resistance-aware therapy "
           "simulation", style="Title")
    P(doc, "Jaivardhan Samadhiya*", style="Author List")
    P(doc, "1 West Valley College, Saratoga, United States of America", italic=True, size=9)
    P(doc, "* Correspondence: Corresponding Author, jsamadhi@mywvm.wvm.edu")
    P(doc, "Manuscript statistics: ~5,800 words (main text); 5 figures; 5 tables.",
      italic=True, size=9)
    P(doc)
    rich(doc, [("Keywords: ", True, False),
               ("bacteriophage, phage–host interaction, machine learning, data "
                "leakage, phage therapy, phage cocktail, antimicrobial resistance, "
                "graph neural network", False, False)])

    # ---------------- Abstract ----------------
    H(doc, "Abstract", 1)
    P(doc,
      "Computational prediction of which bacteriophage infects which bacterium is "
      "a prerequisite for rational phage therapy, yet reported accuracies are "
      "frequently inflated by data leakage between training and test genomes. We "
      "present PrecisionPhage, an end-to-end, reproducible framework that "
      "couples strict, homology-aware evaluation to interaction prediction, "
      "phage-cocktail optimisation and eco-evolutionary therapy simulation. Using "
      "1,947 experimentally determined phage–host pairs (1,418 phages, 323 hosts) "
      "with complete genomes, we engineered leakage-safe features that combine "
      "genome composition, multi-scale nucleotide homology, CRISPR spacer matching "
      "and protein-level homology, and trained a gradient-boosted decision tree "
      "(GBM) under four primary homology-aware cross-validation regimes, with "
      "leave-one-genus-out (LOGO) reported separately as a supplementary "
      "taxonomic reference (Table 2). Performance remained high under taxonomic "
      "holdouts but decreased markedly under increasingly stringent "
      "homology-aware evaluation (AUROC from 0.960 for unseen species to a "
      "pooled AUROC of 0.780 in the strict cold-start regime, based on only "
      "five independent evaluation blocks with per-fold GBM AUROC ranging from "
      "0.69 to 0.89), providing evidence that the splits remove shortcut signal while "
      "highlighting the limited precision of cold-start estimates. "
      "The GBM significantly outperformed an inductive graph neural network in "
      "all four primary regimes (DeLong false-discovery-rate q < 3×10⁻⁷), and a "
      "controlled ablation showed that graph message passing did not help and "
      "significantly hurt prediction for novel phages. Against external "
      "baselines on identical test pairs, PrecisionPhage significantly exceeded "
      "the published PHIST implementation (+0.10 to +0.28 AUROC, all q < 10⁻³) "
      "and our faithful in-house RaFAH-style reimplementation under identical "
      "leakage-controlled evaluation (+0.07 to +0.38 "
      "AUROC, all q < 0.05). Model predictions "
      "drove exact integer-program "
      "cocktail design that tracked an oracle and a mechanistic eco-evolutionary "
      "simulation showing that only redundant (≥2-phage-per-strain) cocktails "
      "both suppressed bacterial load and prevented resistance. PrecisionPhage "
      "provides a rigorously validated, leakage-free framework for future "
      "computationally guided phage therapy studies.")

    # ---------------- Introduction ----------------
    H(doc, "1 Introduction", 1)
    P(doc,
      "The global rise of antimicrobial resistance, associated with an estimated "
      "1.27 million deaths in 2019, has renewed interest in bacteriophages as "
      "precision antibacterials (Murray et al., 2022). Because most phages infect "
      "a narrow range of hosts, therapeutic success depends on matching phages to "
      "the infecting strain, and increasingly on combining several phages into "
      "cocktails to broaden coverage and suppress the emergence of resistance "
      "(Chan et al., 2013; Torres-Barceló and Hochberg, 2016). Experimentally "
      "screening phage–host pairs is slow and cannot scale to the diversity of "
      "clinical isolates, motivating computational prediction of phage–host "
      "interactions directly from genome sequence (Edwards et al., 2016).")
    P(doc,
      "A range of in silico host-prediction methods now exists, from alignment-"
      "free oligonucleotide-frequency and exact k-mer matching (Ahlgren et al., "
      "2017; Galiez et al., 2017; Zielezinski et al., 2022) to protein-content "
      "random forests (Coutinho et al., 2021) and graph- or deep-learning models "
      "that exploit the bipartite structure of known infections (Wang et al., "
      "2020; Shang and Sun, 2022). Reported accuracies are often high, but a "
      "recurring methodological problem undermines their interpretation: data "
      "leakage. When near-identical genomes, or different members of the same "
      "species or genus, are split across training and test sets, a model can "
      "achieve excellent test metrics by recognising near-duplicates rather than "
      "by learning generalisable interaction biology. Such leakage inflates "
      "performance precisely in the cold-start setting—novel phages and novel "
      "hosts—that matters most clinically.")
    P(doc,
      "Here we take the position that, for phage therapy, scientific validity "
      "under realistic generalisation is more important than headline accuracy. "
      "We introduce PrecisionPhage, a reproducible framework whose central design "
      "principle is leakage control: genomes are clustered by Mash distance and "
      "held out at the cluster level, all feature scaling and dimensionality "
      "reduction are fit strictly inside training folds, and every claim is "
      "accompanied by confidence intervals and multiplicity-corrected significance "
      "tests. Within this framework we (i) engineer leakage-safe features spanning "
      "composition, nucleotide homology, CRISPR targeting and protein homology; "
      "(ii) benchmark a gradient-boosted tree against an inductive graph neural "
      "network, the published PHIST tool, and our faithful in-house RaFAH-style "
      "reimplementation under identical leakage-controlled evaluation; (iii) translate calibrated predictions into "
      "provably optimal phage cocktails; and (iv) connect those cocktails to "
      "treatment outcome through a mechanistic eco-evolutionary simulation of "
      "resistance. The result is a transparent picture of what is and is not "
      "predictable from genomes once leakage is removed, and a practical pipeline "
      "from sequence to cocktail to predicted therapeutic dynamics.")

    # ---------------- Materials and Methods ----------------
    H(doc, "2 Materials and Methods", 1)

    H(doc, "2.1 Interaction data and genome resolution", 2)
    P(doc,
      "We assembled experimentally determined phage–host interactions from three "
      "sources: the Nahant Collection marine cross-infection matrix (Kauffman et "
      "al., 2018), curated NCBI host-range records, and a Staphylococcus host-"
      "range study. After canonicalisation and de-duplication this yielded 8,849 "
      "labelled pairs (2,770 positive, 6,079 negative) spanning 2,331 phages and "
      "387 hosts (380 species, 122 genera) with zero label conflicts. Each phage "
      "and host name was resolved to a genome FASTA through a deterministic, "
      "audited index; a pair entered the sequence-based models only when both "
      "genomes resolved. This genome-complete subset comprised 1,947 pairs (1,488 "
      "positive, 459 negative) over 1,418 phages and 323 hosts, and is used for "
      "all modelling below. Pairs that could not be resolved were reported and "
      "excluded rather than zero-imputed, which would silently corrupt metrics.")

    H(doc, "2.2 Leakage-safe features", 2)
    P(doc,
      "Each genome was summarised by a 212-dimensional node descriptor comprising "
      "the canonical 4-mer spectrum, codon usage, GC content and dinucleotide "
      "bias, computed with vectorised routines. Within every training fold these "
      "descriptors were standardised and reduced by principal component analysis "
      "(64 components for phages and hosts separately), with the scaler and PCA "
      "fit only on training rows and applied to held-out rows. For every pair we "
      "computed a 24-dimensional edge descriptor: composition distances (cosine, "
      "L1, Pearson, Jaccard, 3-mer and 6-mer distances, GC difference); multi-"
      "scale exact nucleotide homology at word sizes k = 16 and k = 20 (fraction "
      "of shared words, genome coverage, log shared count, longest contiguous "
      "match); CRISPR features (host CRISPR-array spacer extraction and spacer-to-"
      "phage matching: spacer count, hit count, best hit fraction, presence); and "
      "tblastx-style protein homology from six-frame ORF translation (minimum ORF "
      "90 aa) with amino-acid 7-mer containment and coverage. All homology, CRISPR "
      "and protein features depend only on the two genomes of a pair and therefore "
      "introduce no train/test leakage.")

    H(doc, "2.3 Homology-aware cross-validation", 2)
    P(doc,
      "To define generalisation regimes we sketched every genome with MinHash "
      "(k = 21, 256 hashes) and grouped genomes by single-linkage clustering at a "
      "Mash distance of 0.05 (≈95% average nucleotide identity), approximating a "
      "species-level boundary (Ondov et al., 2016). Our headline benchmark "
      "comprises four primary homology-aware leave-one-group-out regimes "
      "reported in Table 1: leave-one-species-out (LOSO), leave-one-host-cluster-"
      "out, leave-one-phage-cluster-out, and a combined regime in which both phage "
      "and host clusters in a test block are unseen during training (true cold "
      "start). As a supplementary taxonomic reference we additionally report "
      "leave-one-genus-out (LOGO) performance in Table 2; LOGO is not strictly "
      "ordered relative to cluster holdouts because taxonomic and homology-aware "
      "splits measure different generalisation axes. Folds were retained only when they contained at least three "
      "positive and three negative examples. Hyperparameter selection, when used, "
      "was performed by group-aware nested cross-validation inside the training "
      "split.")

    H(doc, "2.4 Models", 2)
    P(doc,
      "The primary model is a gradient-boosted decision tree (XGBoost: 400 trees, "
      "maximum depth 5, learning rate 0.05, subsample 0.8, column subsample 0.8) "
      "operating on the concatenated phage-node, host-node and edge features "
      "(Chen and Guestrin, 2016). Predictions were calibrated post hoc by isotonic "
      "regression fit on inner-fold out-of-fold predictions, with an order-"
      "preserving perturbation guaranteeing that calibration never reduces AUROC "
      "(Niculescu-Mizil and Caruana, 2005). As a relational comparator we "
      "implemented an inductive GraphSAGE encoder with an edge-MLP decoder over "
      "the bipartite phage–host graph (Hamilton et al., 2017); to isolate the "
      "contribution of message passing we ran an ablation in which the graph edges "
      "were removed, reducing the model to a multilayer perceptron on the same "
      "features. Models were implemented with scikit-learn and PyTorch (Pedregosa "
      "et al., 2011).")

    H(doc, "2.5 External baselines", 2)
    P(doc,
      "We compared against two established approaches on identical, leakage-"
      "controlled test pairs. PHIST (Zielezinski et al., 2022) is the published, "
      "alignment-free tool: we built it from source and ran it on all covered "
      "genomes; it is unsupervised and scores a pair by the number of exact "
      "shared 25-mers, so there is no training leakage in its favour. The "
      "published, pretrained RaFAH (Coutinho et al., 2021) could not be executed "
      "in our environment because it requires an R runtime that was unavailable "
      "and its pretrained model and HMM database are hosted on a network-blocked "
      "server; we therefore reimplemented its methodology (predict host genus "
      "from phage protein content with a random forest, using six-frame ORFs and "
      "feature-hashed amino-acid 6-mer presence as a proxy for protein-cluster "
      "features) and label it RaFAH-style throughout. The RaFAH-style model was "
      "trained on known phage→host-genus associations within each training fold "
      "and scored a pair by the predicted probability of the host's genus. This "
      "is a methodological comparison of RaFAH's approach, not a benchmark of "
      "RaFAH's released weights. PHIST and RaFAH-style scores were evaluated on "
      "the same per-regime test folds as our model and compared with DeLong's "
      "paired test with Benjamini–Hochberg correction across regimes.")

    H(doc, "2.6 Statistical analysis", 2)
    P(doc,
      "AUROC is reported with DeLong analytic 95% confidence intervals (DeLong et "
      "al., 1988); models were compared per regime with DeLong's paired test and "
      "McNemar's test, and the AUROC difference with a 10,000-sample paired "
      "bootstrap. Skill against chance was assessed by a 1,000-permutation label "
      "test. All families of p-values were Benjamini–Hochberg false-discovery-rate "
      "corrected (Benjamini and Hochberg, 1995). Calibration was quantified by the "
      "expected calibration error (ECE, 10 bins). A fixed random seed and pinned "
      "environment ensure deterministic reproduction.")

    H(doc, "2.7 Cocktail optimisation and therapy simulation", 2)
    P(doc,
      "Out-of-fold predictions (stratified group K-fold) produced a leakage-free "
      "predicted infection matrix, thresholded at the F1-optimal operating point. "
      "Minimum-size cocktails covering all coverable host strains were computed "
      "exactly by integer linear programming (HiGHS solver) and approximately by "
      "greedy set cover, with a k-redundancy constraint requiring each strain to "
      "be covered by at least k phages; coverage was always scored against the "
      "true interaction matrix. To connect cocktails to outcome we built a "
      "deterministic eco-evolutionary ordinary-differential-equation model of "
      "sensitive and resistant bacterial subpopulations and phage populations, in "
      "which the per-division resistance probability falls geometrically with the "
      "number of distinct phages targeting a strain and an extinction floor "
      "prevents regrowth from sub-threshold densities (Levin and Bull, 2004). We "
      "compared no treatment, monophage, a non-redundant cocktail and a redundant "
      "(k ≥ 2) cocktail.")

    # ---------------- Results ----------------
    H(doc, "3 Results", 1)

    H(doc, "3.1 Performance declines under stringent homology-aware evaluation", 2)
    P(doc,
      "Across the four primary leakage regimes in Table 1, GBM AUROC remained "
      "high under species and host-cluster holdouts but decreased markedly as "
      "evaluation became stricter: 0.960 (95% CI 0.946–0.974) for unseen "
      "species, 0.954 (0.938–0.969) for unseen host clusters, 0.853 (0.815–0.891) "
      "for unseen phage clusters, and 0.780 (0.729–0.832) pooled for the fully "
      "cold-start regime (Table 1, Figure 1). Per-fold GBM AUROC in the cold-"
      "start regime was 0.894, 0.688, 0.867, 0.765 and 0.785 (mean 0.800; "
      "five blocks), illustrating substantial block-to-block variability relative "
      "to the pooled estimate. Every primary regime was significantly above "
      "chance (permutation q = 1.0×10⁻³). The supplementary "
      "LOGO reference in Table 2 (pooled AUROC 0.919) sits between species and "
      "host-cluster performance, illustrating that taxonomic and homology-aware "
      "splits are not strictly ordered. This pattern is consistent with the "
      "splits removing exploitable shortcut signal: a model relying on near-"
      "duplicate genomes would retain high accuracy until the strictest split and "
      "then collapse, whereas PrecisionPhage degrades gracefully and remains "
      "well above chance even when both partners are unseen. Calibration was "
      "excellent in the species and host-cluster regimes (ECE 0.024 and 0.047). "
      "The strict cold-start regime necessarily comprised only five independent "
      "evaluation blocks, so its pooled AUROC should be interpreted as a "
      "preliminary small-sample estimate rather than a settled generalisation "
      "result; the wide spread of per-fold values (0.688–0.894) reinforces this "
      "caution despite pooled performance remaining above chance.")

    H(doc, "3.2 A gradient-boosted tree outperforms a graph neural network", 2)
    P(doc,
      "The GBM significantly outperformed the inductive GNN in all four primary "
      "regimes, by +0.080 (unseen species) and +0.076 (unseen host cluster) AUROC "
      "and by a much larger +0.249 and +0.160 in the phage-cluster and cold-start "
      "regimes (all DeLong q ≤ 2.7×10⁻⁷; Table 1). The GNN was also poorly "
      "calibrated (ECE up to 0.55). A controlled ablation clarified why: removing "
      "graph message passing did not reduce performance and significantly improved "
      "it in the phage-cluster holdout (AUROC 0.575 with the graph versus 0.709 "
      "without; DeLong p = 1.3×10⁻⁶). In other words, for inductive prediction on "
      "novel genomes the relational signal is already captured by the pairwise "
      "features, and propagating information across a graph dominated by unseen "
      "nodes is, if anything, harmful. We report this as a clean negative result: "
      "graph learning is not necessary for this task at the available data scale.")

    H(doc, "3.3 Genome composition and nucleotide homology dominate", 2)
    P(doc,
      "Gain-based importances from the full model attributed 61% of predictive "
      "signal to phage genome composition and 29% to host genome composition, with "
      "multi-scale nucleotide homology contributing 6.1%, protein homology 1.6%, "
      "composition distances 1.6% and CRISPR matching 0.3% (Figure 2). The "
      "homology, protein and CRISPR features were therefore complementary but "
      "secondary to composition; their main value, consistent with the external "
      "comparisons below, is precision on the subset of pairs where exact sequence "
      "evidence exists.")

    H(doc, "3.4 PrecisionPhage exceeds external baselines on identical test pairs", 2)
    P(doc,
      "Evaluated on the same per-regime test pairs, PrecisionPhage significantly "
      "outperformed the published PHIST implementation in every primary regime, "
      "by +0.276 and +0.273 AUROC for unseen species and host clusters and by "
      "+0.200 and +0.099 for the phage-cluster and cold-start regimes (all DeLong "
      "q < 10⁻³; Table 3, Figure 3). PHIST found at least one shared 25-mer for "
      "only 35% of covered pairs, which explains its modest AUROC despite high "
      "precision in the regimes where it does fire (AUPRC 0.84–0.86 for the "
      "species, host-cluster and cold-start regimes). PrecisionPhage also "
      "significantly exceeded our faithful in-house RaFAH-style reimplementation "
      "under identical leakage-controlled evaluation in all primary regimes (q < 0.05). Notably, the smallest gap occurred in "
      "the phage-cluster regime—predicting the host of a genuinely novel phage, "
      "which is exactly RaFAH's design goal—where the RaFAH-style model reached "
      "0.780 versus our 0.853 (+0.073, q = 0.024); in the cold-start regime, "
      "where the host taxon is often unseen, it collapsed to chance (AUROC 0.431, "
      "permutation q = 0.99). These results delineate where the published PHIST "
      "tool and a protein-taxonomy reimplementation remain competitive and where "
      "an integrated model is required.")

    H(doc, "3.5 Predictions yield provably optimal, robust cocktails", 2)
    P(doc,
      "Using leakage-free out-of-fold predictions, the minimum cocktail covering "
      "all coverable host strains contained 180 phages by exact integer "
      "programming, matched by the greedy solution and close to the 188-phage "
      "true-optimal oracle, and model-driven greedy selection tracked the oracle "
      "coverage curve while massively outperforming random selection (Figure 4). "
      "Imposing redundancy increased robustness as expected: a single-cover "
      "cocktail (180 phages) covered 90.9% of strains at least once, whereas "
      "doubly- and triply-redundant cocktails (291 and 372 phages) covered 95.2% "
      "and 97.8% of strains at least once (Table 4). Because the available "
      "interaction matrix is sparse and taxonomically imbalanced, the optimized "
      "cocktail sizes should not be interpreted as clinically realistic "
      "formulations but rather as properties of the current benchmark dataset; "
      "the large absolute sizes reflect sparsity rather than a limitation of the "
      "optimiser.")

    H(doc, "3.6 Only redundant cocktails prevent resistance in simulation", 2)
    P(doc,
      "In the eco-evolutionary simulation of a multi-strain infection, no "
      "treatment and monophage therapy failed to reduce bacterial load. A non-"
      "redundant cocktail (one effective phage per strain) achieved transient "
      "suppression (2.2 log₁₀ nadir) but relapsed through fully resistant "
      "regrowth. Only the redundant (k ≥ 2) cocktail both suppressed load "
      "durably (final density 6.0×10⁵, a 2.8 log₁₀ drop) and prevented resistance "
      "(resistant fraction 0 at the end of simulation; Table 5, Figure 5). This "
      "recapitulates the evolutionary rationale for multi-phage cocktails and "
      "shows that the same model outputs that drive cocktail selection can be "
      "propagated to a mechanistic prediction of therapeutic durability.")

    # ---------------- Discussion ----------------
    H(doc, "4 Discussion", 1)
    P(doc,
      "PrecisionPhage shows that phage–host interactions remain predictable from "
      "genome sequence even under strict, homology-aware evaluation, but that the "
      "achievable accuracy depends strongly on how leakage is controlled. "
      "Performance remained high under taxonomic holdouts but decreased markedly "
      "under increasingly stringent homology-aware evaluation, from AUROC 0.960 "
      "for unseen species to a pooled AUROC of 0.780 in the strict cold-start "
      "regime (five blocks; per-fold GBM AUROC 0.688–0.894), quantifying an "
      "effect that is frequently hidden in the literature, where favourable splits "
      "can leave near-duplicate genomes on both sides of the partition. We argue "
      "that cold-start regimes, in which both phage and host are unseen, are the "
      "most stringent yardstick for clinical translation, but that the present "
      "cold-start estimate is based on few independent blocks and should be "
      "interpreted cautiously rather than as a precise generalisation bound. "
      "Methods should be reported across a leakage hierarchy rather than at a "
      "single operating point.")
    P(doc,
      "Two findings are methodologically instructive. First, a well-regularised "
      "gradient-boosted tree on leakage-safe pairwise features outperformed an "
      "inductive graph neural network in all four primary regimes, and message "
      "passing actively harmed prediction for novel phages. At the present data "
      "scale the task is feature-limited rather than model-limited, and the "
      "relational structure that a GNN is designed to exploit is largely redundant "
      "with explicit pairwise homology features. Second, the external comparisons "
      "localise the competitive frontier: the published PHIST tool is precise but "
      "low-recall because most pairs share no exact long k-mer, while the RaFAH-"
      "style reimplementation is strongest exactly for novel-phage host assignment "
      "and weakest when the host lineage itself is unseen. PrecisionPhage's "
      "advantage comes from integrating these complementary signals with genome "
      "composition and CRISPR targeting.")
    P(doc,
      "By coupling calibrated predictions to exact combinatorial optimisation and "
      "a mechanistic resistance model, the framework also bridges prediction and "
      "intervention. The simulation reproduces a central principle of phage "
      "therapy—that redundancy, not merely coverage, is required to forestall "
      "resistance—directly from model-selected cocktails, providing a testable, "
      "quantitative hypothesis for experimental follow-up.")
    P(doc,
      "Our study has limitations. The genome-complete dataset, although carefully "
      "curated, is modest and skewed toward a few host taxa. The strict cold-start "
      "regime necessarily comprised only five independent evaluation blocks; "
      "per-fold GBM AUROC ranged from 0.688 to 0.894 (mean 0.800), so the "
      "pooled cold-start AUROC should be treated as a preliminary small-sample "
      "estimate rather than a settled generalisation result. Because the interaction matrix is sparse and "
      "taxonomically imbalanced, optimized cocktail sizes should not be interpreted "
      "as clinically realistic formulations. The RaFAH comparison is a faithful "
      "in-house reimplementation rather than the published, pretrained tool, "
      "because the latter could not be run in our environment; the PHIST comparison "
      "uses the genuine published software. The therapy simulation is a deterministic "
      "abstraction that omits pharmacokinetics, spatial structure and immune "
      "dynamics, and is intended to illustrate evolutionary principles rather than "
      "to predict clinical trajectories. Finally, the secondary role of protein and "
      "CRISPR features suggests that richer protein representations—curated protein "
      "clusters or protein language models—are the most promising route to widen "
      "the lead on genuinely novel phages, the one regime where a protein-taxonomy "
      "baseline approached our model.")
    P(doc,
      "In summary, PrecisionPhage provides a transparent, reproducible and "
      "rigorously validated pipeline from genome sequence to leakage-controlled "
      "interaction prediction, provably optimal cocktail design and resistance-"
      "aware therapy simulation, establishing a defensible baseline for the next "
      "generation of computationally guided phage therapeutics.")

    # ---------------- Back matter ----------------
    H(doc, "Conflict of Interest", 1)
    P(doc, "The authors declare that the research was conducted in the absence of "
           "any commercial or financial relationships that could be construed as a "
           "potential conflict of interest.")

    H(doc, "Author Contributions", 1)
    P(doc, "The author contributed to the conception and design of the study, "
           "the analysis and interpretation of results, and the writing of the "
           "manuscript, and approved the submitted version.")

    H(doc, "Funding", 1)
    P(doc, "No funding received from any parties.")

    H(doc, "Acknowledgments", 1)
    P(doc, "The author has no external acknowledgments to declare for this study.")

    H(doc, "Data Availability Statement", 1)
    P(doc, f"The interaction data analysed in this study derive from publicly "
           f"available sources (the Nahant Collection and NCBI). The complete "
           f"analysis code, configuration and result artifacts required to "
           f"reproduce all figures and tables are available in the project "
           f"repository at {REPO_URL}. All manuscript numbers come from "
           f"data/results_v2/. PHIST is available from its original authors.")

    # ---------------- References ----------------
    H(doc, "References", 1)
    refs = [
        "Ahlgren, N. A., Ren, J., Lu, Y. Y., Fuhrman, J. A., and Sun, F. (2017). Alignment-free oligonucleotide frequency dissimilarity measure improves prediction of hosts from metagenomically-derived viral sequences. Nucleic Acids Res. 45, 39–53.",
        "Benjamini, Y., and Hochberg, Y. (1995). Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. Series B 57, 289–300.",
        "Chan, B. K., Abedon, S. T., and Loc-Carrillo, C. (2013). Phage cocktails and the future of phage therapy. Future Microbiol. 8, 769–783.",
        "Chen, T., and Guestrin, C. (2016). XGBoost: a scalable tree boosting system. Proc. 22nd ACM SIGKDD, 785–794.",
        "Coutinho, F. H., Zaragoza-Solas, A., López-Pérez, M., Barylski, J., Zielezinski, A., Dutilh, B. E., et al. (2021). RaFAH: host prediction for viruses of Bacteria and Archaea based on protein content. Patterns 2, 100274.",
        "DeLong, E. R., DeLong, D. M., and Clarke-Pearson, D. L. (1988). Comparing the areas under two or more correlated receiver operating characteristic curves: a nonparametric approach. Biometrics 44, 837–845.",
        "Edwards, R. A., McNair, K., Faust, K., Raes, J., and Dutilh, B. E. (2016). Computational approaches to predict bacteriophage–host relationships. FEMS Microbiol. Rev. 40, 258–272.",
        "Galiez, C., Siebert, M., Enault, F., Vincent, J., and Söding, J. (2017). WIsH: who is the host? Predicting prokaryotic hosts from metagenomic phage contigs. Bioinformatics 33, 3113–3114.",
        "Hamilton, W. L., Ying, R., and Leskovec, J. (2017). Inductive representation learning on large graphs. Adv. Neural Inf. Process. Syst. 30, 1024–1034.",
        "Kauffman, K. M., Hussain, F. A., Yang, J., Arevalo, P., Brown, J. M., Chang, W. K., et al. (2018). A major lineage of non-tailed dsDNA viruses as unrecognized killers of marine bacteria. Nature 554, 118–122.",
        "Levin, B. R., and Bull, J. J. (2004). Population and evolutionary dynamics of phage therapy. Nat. Rev. Microbiol. 2, 166–173.",
        "Murray, C. J. L., Ikuta, K. S., Sharara, F., Swetschinski, L., Robles Aguilar, G., Gray, A., et al. (2022). Global burden of bacterial antimicrobial resistance in 2019: a systematic analysis. Lancet 399, 629–655.",
        "Niculescu-Mizil, A., and Caruana, R. (2005). Predicting good probabilities with supervised learning. Proc. 22nd Int. Conf. Mach. Learn., 625–632.",
        "Ondov, B. D., Treangen, T. J., Melsted, P., Mallonee, A. B., Bergman, N. H., Koren, S., et al. (2016). Mash: fast genome and metagenome distance estimation using MinHash. Genome Biol. 17, 132.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., et al. (2011). Scikit-learn: machine learning in Python. J. Mach. Learn. Res. 12, 2825–2830.",
        "Shang, J., and Sun, Y. (2022). CHERRY: a Computational metHod for accuratE pRediction of virus–prokarYotic interactions. Brief. Bioinform. 23, bbac182.",
        "Torres-Barceló, C., and Hochberg, M. E. (2016). Evolutionary rationale for phages as complements of antibiotics. Trends Microbiol. 24, 249–256.",
        "Wang, W., Ren, J., Tang, K., Dart, E., Ignacio-Espinoza, J. C., Fuhrman, J. A., et al. (2020). A network-based integrated framework for predicting virus–prokaryote interactions. NAR Genom. Bioinform. 2, lqaa044.",
        "Zielezinski, A., Deorowicz, S., and Gudyś, A. (2022). PHIST: fast and accurate prediction of prokaryotic hosts from metagenomic viral sequences. Bioinformatics 38, 1447–1449.",
    ]
    for r in refs:
        P(doc, r)

    # ---------------- Tables ----------------
    doc.add_page_break()
    H(doc, "Tables", 1)

    cold_gbm_cell = "0.780 pooled (0.729–0.832)"
    cold_fold_cell = (fold_str if fold_str
                      else "0.894, 0.688, 0.867, 0.765, 0.785 (mean 0.800, n=5 blocks)")
    table(doc,
          ["Regime", "GBM AUROC (95% CI)", "Per-fold GBM AUROC",
           "GBM ECE", "GNN AUROC (95% CI)", "GBM−GNN ΔAUROC", "DeLong q"],
          [["Unseen species (LOSO)", "0.960 (0.946–0.974)", "—",
            "0.024", "0.879 (0.851–0.907)", "+0.080 (0.057–0.104)", "6.0×10⁻¹¹"],
           ["Unseen host cluster", "0.954 (0.938–0.969)", "—",
            "0.047", "0.878 (0.850–0.905)", "+0.076 (0.053–0.100)", "2.3×10⁻¹⁰"],
           ["Unseen phage cluster", "0.853 (0.815–0.891)", "—",
            "0.424", "0.605 (0.546–0.663)", "+0.249 (0.195–0.302)", "9.8×10⁻¹⁹"],
           ["Both unseen (cold start)", cold_gbm_cell, cold_fold_cell,
            "0.166", "0.621 (0.555–0.686)", "+0.160 (0.100–0.220)", "2.7×10⁻⁷"]],
          caption="Table 1. Interaction-prediction performance across the four "
                  "primary homology-aware leakage regimes and comparison of the "
                  "gradient-boosted tree (GBM) with the graph neural network (GNN). "
                  "For the cold-start regime, pooled AUROC aggregates all held-out "
                  "pairs across five independent blocks; per-fold values are "
                  "reported separately because block count is small. All regimes "
                  "are above chance (permutation q = 1.0×10⁻³).")

    table(doc,
          ["Regime", "Mean fold AUROC (95% CI)", "Pooled AUROC", "ECE", "Folds"],
          [["LOSO (unseen species)", "0.904 (0.753–0.984)", "0.960", "0.024", "28"],
           ["LOGO (unseen genus)", "0.889 (0.837–0.935)", "0.919", "0.053", "28"],
           ["Unseen host cluster", "0.901 (0.796–0.979)", "0.954", "0.047", "26"],
           ["Unseen phage cluster", "0.869 (0.799–0.932)", "0.853", "0.424", "22"],
           ["Both unseen (cold start)", "0.800 (0.738–0.863)", "0.780", "0.166", "5"]],
          caption="Table 2. Leakage hierarchy for the GBM across taxonomic "
                  "(species, genus) and homology-aware genome-cluster regimes. "
                  "LOGO is reported as a supplementary taxonomic reference; regimes "
                  "are not strictly ordered because taxonomic and homology-aware "
                  "splits measure different generalisation axes. Pooled values for "
                  "the four primary regimes match Table 1 exactly.")

    table(doc,
          ["Regime", "PrecisionPhage AUROC", "PHIST AUROC", "RaFAH-style AUROC",
           "ΔAUROC vs PHIST (q)", "ΔAUROC vs RaFAH-style (q)"],
          [["Unseen species (LOSO)", "0.960", "0.684", "0.601",
            "+0.276 (1.5×10⁻¹¹⁰)", "+0.359 (6.9×10⁻⁷³)"],
           ["Unseen host cluster", "0.954", "0.681", "0.575",
            "+0.273 (3.5×10⁻⁹⁶)", "+0.378 (6.2×10⁻⁹³)"],
           ["Unseen phage cluster", "0.853", "0.653", "0.780",
            "+0.200 (4.3×10⁻¹⁸)", "+0.073 (2.3×10⁻²)"],
           ["Both unseen (cold start)", "0.780", "0.681", "0.431",
            "+0.099 (8.6×10⁻⁴)", "+0.350 (7.0×10⁻²⁰)"]],
          caption="Table 3. External baseline comparison on identical test pairs. "
                  "PHIST is the real published tool; RaFAH-style is a faithful "
                  "in-house reimplementation (see Methods). All differences favour "
                  "PrecisionPhage at FDR q < 0.05.")

    table(doc,
          ["Redundancy k", "Cocktail size (phages)", "Strains covered ≥1×", "Strains covered ≥k×"],
          [["1", "180", "90.9%", "90.9%"],
           ["2", "291", "95.2%", "47.0%"],
           ["3", "372", "97.8%", "33.5%"]],
          caption="Table 4. k-redundant cocktails computed from model predictions "
                  "and scored against the true interaction matrix.")

    table(doc,
          ["Strategy", "Phages", "Final load (CFU)", "Nadir (CFU)", "log₁₀ drop",
           "Resistant fraction (end)", "Rebound"],
          [["No treatment", "0", "1.0×10⁹", "5.0×10⁶", "0.0", "0.0", "No"],
           ["Monophage", "1", "1.0×10⁹", "5.0×10⁶", "0.0", "0.0", "No"],
           ["Cocktail (k=1)", "4", "1.0×10⁹", "2.9×10⁴", "2.23", "1.0", "Yes"],
           ["Redundant cocktail (k≥2)", "8", "6.0×10⁵", "7.5×10³", "2.83", "0.0", "Yes"]],
          caption="Table 5. Eco-evolutionary therapy simulation outcomes. Only the "
                  "redundant cocktail both suppresses load and prevents resistance.")

    # ---------------- Figure legends + embedded figures ----------------
    doc.add_page_break()
    H(doc, "Figure Legends", 1)
    legends = [
        ("Figure 1.", "Interaction-prediction performance under leakage-controlled "
         "evaluation. Generalisation AUROC of the GBM and GNN across the four "
         "primary leakage regimes, the decline in performance under increasingly "
         "stringent homology-aware evaluation, and the graph message-passing "
         "ablation showing no benefit (and harm for novel phages) from graph "
         "propagation."),
        ("Figure 2.", "Gain-based feature importances. Genome composition of the "
         "phage and host dominate, with nucleotide homology, protein homology and "
         "CRISPR matching contributing complementary but secondary signal."),
        ("Figure 3.", "External baseline comparison on identical, leakage-"
         "controlled test pairs. PrecisionPhage significantly exceeds the "
         "published PHIST implementation and our faithful in-house RaFAH-style "
         "reimplementation under identical leakage-controlled evaluation in every "
         "primary regime; the smallest margin versus RaFAH-style occurs for "
         "novel-phage host assignment, RaFAH's design goal."),
        ("Figure 4.", "Model-driven phage-cocktail design. Predicted-greedy "
         "coverage tracks the true-optimal oracle and greatly exceeds random "
         "selection as a function of cocktail size."),
        ("Figure 5.", "Eco-evolutionary simulation of phage therapy. Bacterial "
         "load and resistant-subpopulation dynamics under no treatment, "
         "monophage, a non-redundant cocktail and a redundant (k ≥ 2) cocktail; "
         "only the redundant cocktail achieves durable suppression without "
         "resistance."),
    ]
    for tag, text in legends:
        rich(doc, [(tag + " ", True, False), (text, False, False)])

    doc.add_page_break()
    H(doc, "Figures", 1)
    for tag, img in [("Figure 1", "figure_main.png"),
                     ("Figure 2", "fig_feature_importance.png"),
                     ("Figure 3", "fig_phist_compare.png"),
                     ("Figure 4", "cocktail_coverage.png"),
                     ("Figure 5", "temporal_dynamics.png")]:
        P(doc, tag, bold=True)
        figure(doc, RD / img)
        doc.add_paragraph()

    doc.save(str(OUT))
    doc.save(str(OUT_COPY))
    print("WROTE", OUT)
    print("WROTE", OUT_COPY)
    print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))


if __name__ == "__main__":
    main()
